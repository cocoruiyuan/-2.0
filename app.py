from __future__ import annotations

from dataclasses import dataclass, replace
from io import BytesIO
from pathlib import Path
import math
import random
import zipfile

import fitz  # PyMuPDF
import streamlit as st
from docx import Document
from PIL import (
    Image,
    ImageChops,
    ImageDraw,
    ImageEnhance,
    ImageFilter,
    ImageFont,
    ImageOps,
)


# ============================================================
# 页面设置
# ============================================================
PAGE_WIDTH = 1240
PAGE_HEIGHT = 1754

LEFT_MARGIN = 110
RIGHT_MARGIN = 80
TOP_MARGIN = 125
BOTTOM_MARGIN = 100

BASE_DIR = Path(__file__).resolve().parent

SUPPORTED_FILE_TYPES = ["pdf", "docx", "txt"]


@dataclass(frozen=True)
class RenderSettings:
    font_path: str
    font_size: int
    line_spacing: int
    paper_type: str
    paper_age: float
    wrinkle_strength: float
    photo_effect: str
    ink_name: str
    pen_style: str
    texture_strength: float
    font_weight: float
    randomness: int
    seed: int
    slant: float
    word_spacing: int
    connection_strength: float
    baseline_wave: float
    flourish_level: int
    correction_level: int
    correction_style: str
    teacher_marks: int
    header_enabled: bool
    header_date: str
    header_lesson: str
    show_page_number: bool
    quality_scale: int = 3

    @property
    def rule_spacing(self) -> int:
        return max(
            self.font_size + 8,
            self.font_size + self.line_spacing,
        )



PRESETS: dict[str, dict[str, object]] = {
    "参考图：蓝色圆珠笔": {
        "font_size": 49,
        "line_spacing": 10,
        "paper_type": "俄语作业本",
        "paper_age": 0.12,
        "wrinkle_strength": 0.10,
        "photo_effect": "无",
        "ink_name": "蓝色",
        "pen_style": "圆珠笔",
        "texture_strength": 0.38,
        "randomness": 1,
        "slant": 0.10,
        "connection_strength": 0.82,
        "word_spacing": -1,
        "baseline_wave": 0.8,
        "flourish_level": 1,
        "correction_level": 1,
        "teacher_marks": 0,
        "correction_style": "混合",
    },
    "整齐钢笔作业": {
        "font_size": 48,
        "line_spacing": 10,
        "paper_type": "俄语作业本",
        "paper_age": 0.05,
        "wrinkle_strength": 0.04,
        "photo_effect": "无",
        "ink_name": "蓝色",
        "pen_style": "钢笔",
        "texture_strength": 0.24,
        "randomness": 1,
        "slant": 0.11,
        "connection_strength": 0.86,
        "word_spacing": -1,
        "baseline_wave": 0.5,
        "flourish_level": 2,
        "correction_level": 0,
        "teacher_marks": 0,
        "correction_style": "混合",
    },
    "黑色中性笔作业": {
        "font_size": 48,
        "line_spacing": 10,
        "paper_type": "普通横线本",
        "paper_age": 0.08,
        "wrinkle_strength": 0.06,
        "photo_effect": "无",
        "ink_name": "黑色",
        "pen_style": "中性笔",
        "texture_strength": 0.22,
        "randomness": 1,
        "slant": 0.08,
        "connection_strength": 0.76,
        "word_spacing": 0,
        "baseline_wave": 0.7,
        "flourish_level": 1,
        "correction_level": 1,
        "teacher_marks": 0,
        "correction_style": "混合",
    },
    "铅笔课堂草稿": {
        "font_size": 47,
        "line_spacing": 9,
        "paper_type": "方格本",
        "paper_age": 0.18,
        "wrinkle_strength": 0.14,
        "photo_effect": "扫描件",
        "ink_name": "黑色",
        "pen_style": "铅笔",
        "texture_strength": 0.68,
        "randomness": 2,
        "slant": 0.06,
        "connection_strength": 0.64,
        "word_spacing": 0,
        "baseline_wave": 1.4,
        "flourish_level": 0,
        "correction_level": 2,
        "teacher_marks": 0,
        "correction_style": "混合",
    },
    "凌乱学生课堂笔记": {
        "font_size": 50,
        "line_spacing": 10,
        "paper_type": "普通横线本",
        "paper_age": 0.22,
        "wrinkle_strength": 0.22,
        "photo_effect": "手机拍照",
        "ink_name": "蓝色",
        "pen_style": "圆珠笔",
        "texture_strength": 0.48,
        "randomness": 3,
        "slant": 0.12,
        "connection_strength": 0.78,
        "word_spacing": -2,
        "baseline_wave": 2.0,
        "flourish_level": 2,
        "correction_level": 2,
        "teacher_marks": 0,
        "correction_style": "混合",
    },
    "老师批改后的作业": {
        "font_size": 48,
        "line_spacing": 10,
        "paper_type": "俄语作业本",
        "paper_age": 0.14,
        "wrinkle_strength": 0.12,
        "photo_effect": "手机拍照",
        "ink_name": "蓝色",
        "pen_style": "圆珠笔",
        "texture_strength": 0.38,
        "randomness": 1,
        "slant": 0.10,
        "connection_strength": 0.82,
        "word_spacing": -1,
        "baseline_wave": 0.8,
        "flourish_level": 1,
        "correction_level": 1,
        "teacher_marks": 2,
        "correction_style": "混合",
    },
}


def option_index(
    options: list,
    value,
) -> int:
    try:
        return options.index(value)
    except ValueError:
        return 0


# ============================================================
# 通用工具
# ============================================================
def clamp(value: int) -> int:
    return max(0, min(255, value))


def normalize_text(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n").strip()


def get_font_files() -> dict[str, Path]:
    discovered: list[Path] = []

    for pattern in ("*.ttf", "*.otf", "*.TTF", "*.OTF"):
        discovered.extend(BASE_DIR.rglob(pattern))

    unique_files = {
        path.resolve()
        for path in discovered
        if path.is_file()
        and ".venv" not in path.parts
        and "site-packages" not in path.parts
    }

    files = sorted(
        unique_files,
        key=lambda path: str(path).lower(),
    )

    return {
        str(path.relative_to(BASE_DIR)): path
        for path in files
    }


@st.cache_resource(show_spinner=False)
def load_font(
    font_path: str,
    font_size: int,
) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(font_path, font_size)


def draw_text_with_cursive_features(
    draw: ImageDraw.ImageDraw,
    position: tuple[int | float, int | float],
    text: str,
    font: ImageFont.FreeTypeFont,
    fill,
    anchor: str | None = None,
) -> None:
    """
    尝试启用字体中的连字和上下文替换功能。

    如果服务器上的 Pillow 没有 RAQM 支持，
    会自动退回普通绘制，不会导致程序崩溃。
    """
    kwargs = {
        "font": font,
        "fill": fill,
    }

    if anchor is not None:
        kwargs["anchor"] = anchor

    try:
        draw.text(
            position,
            text,
            features=["calt", "liga", "clig"],
            **kwargs,
        )
    except Exception:
        draw.text(
            position,
            text,
            **kwargs,
        )


def get_pen_profile(
    settings: RenderSettings,
) -> dict[str, float | tuple[int, int, int]]:
    """
    不同笔迹工具的基础参数。
    """
    if settings.pen_style == "钢笔":
        if settings.ink_name == "蓝色":
            base_color = (34, 57, 116)
        else:
            base_color = (28, 28, 30)

        return {
            "base_color": base_color,
            "variation": 4,
            "pressure_low": 236,
            "pressure_high": 255,
            "alpha_low": 235,
            "alpha_high": 250,
            "blur": 0.030,
            "grain": 0.10 + settings.texture_strength * 0.12,
            "skipiness": 0.02,
            "feather": 0.12 + settings.texture_strength * 0.10,
        }

    if settings.pen_style == "圆珠笔":
        if settings.ink_name == "蓝色":
            base_color = (48, 76, 140)
        else:
            base_color = (40, 40, 42)

        return {
            "base_color": base_color,
            "variation": 5,
            "pressure_low": 228,
            "pressure_high": 255,
            "alpha_low": 224,
            "alpha_high": 246,
            "blur": 0.028,
            "grain": 0.14 + settings.texture_strength * 0.12,
            "skipiness": 0.06 + settings.texture_strength * 0.05,
            "feather": 0.06,
        }

    if settings.pen_style == "铅笔":
        return {
            "base_color": (88, 88, 92),
            "variation": 8,
            "pressure_low": 190,
            "pressure_high": 245,
            "alpha_low": 160,
            "alpha_high": 220,
            "blur": 0.020,
            "grain": 0.34 + settings.texture_strength * 0.26,
            "skipiness": 0.16 + settings.texture_strength * 0.12,
            "feather": 0.02,
        }

    # 默认：中性笔
    if settings.ink_name == "蓝色":
        base_color = (36, 61, 122)
    else:
        base_color = (35, 35, 36)

    return {
        "base_color": base_color,
        "variation": 3,
        "pressure_low": 238,
        "pressure_high": 255,
        "alpha_low": 236,
        "alpha_high": 250,
        "blur": 0.026,
        "grain": 0.08 + settings.texture_strength * 0.08,
        "skipiness": 0.02,
        "feather": 0.04,
    }


def varied_ink_color(
    settings: RenderSettings,
    rng: random.Random,
) -> tuple[int, int, int]:
    profile = get_pen_profile(settings)
    base = profile["base_color"]
    variation = int(profile["variation"])

    return (
        clamp(base[0] + rng.randint(-variation, variation)),
        clamp(base[1] + rng.randint(-variation, variation)),
        clamp(base[2] + rng.randint(-variation, variation)),
    )


def create_pressure_mask(
    text_mask: Image.Image,
    settings: RenderSettings,
    rng: random.Random,
) -> Image.Image:
    """
    生成基础笔压透明度。
    """
    profile = get_pen_profile(settings)

    width, height = text_mask.size
    small_width = max(2, width // 35)
    small_height = max(2, height // 14)

    pressure_pixels = bytes(
        rng.randint(
            int(profile["pressure_low"]),
            int(profile["pressure_high"]),
        )
        for _ in range(small_width * small_height)
    )

    pressure = Image.frombytes(
        "L",
        (small_width, small_height),
        pressure_pixels,
    ).resize(
        (width, height),
        Image.Resampling.BILINEAR,
    )

    return ImageChops.multiply(text_mask, pressure)


def create_stroke_texture_mask(
    size: tuple[int, int],
    settings: RenderSettings,
    rng: random.Random,
) -> Image.Image:
    """
    为笔迹添加纹理：
    - 钢笔：比较顺滑，只带轻微墨水不均
    - 圆珠笔：有轻微断墨和粗糙感
    - 中性笔：较稳定
    - 铅笔：明显颗粒感和石墨摩擦痕迹
    """
    width, height = size
    profile = get_pen_profile(settings)
    grain = float(profile["grain"])
    skipiness = float(profile["skipiness"])

    small_width = max(2, width // 10)
    small_height = max(2, height // 6)

    pixels = bytearray()

    for _ in range(small_width * small_height):
        base_value = 255 - int(rng.random() * 255 * grain * 0.55)
        if rng.random() < skipiness:
            base_value -= rng.randint(12, 55)
        pixels.append(clamp(base_value))

    texture = Image.frombytes(
        "L",
        (small_width, small_height),
        bytes(pixels),
    ).resize(
        (width, height),
        Image.Resampling.BILINEAR,
    )

    if settings.pen_style == "铅笔":
        # 再叠加一层水平方向的石墨纹理
        graphite = Image.new("L", (width, height), 255)
        gdraw = ImageDraw.Draw(graphite)

        line_count = max(6, height // 4)
        for _ in range(line_count):
            y = rng.randint(0, height - 1)
            gdraw.line(
                (
                    0,
                    y,
                    width,
                    y + rng.randint(-1, 1),
                ),
                fill=clamp(232 - rng.randint(0, 42)),
                width=1,
            )

        texture = ImageChops.multiply(texture, graphite)

    return texture


def apply_pen_texture(
    text_mask: Image.Image,
    settings: RenderSettings,
    rng: random.Random,
) -> tuple[Image.Image, tuple[int, int, int], int]:
    """
    根据书写工具把基础文字 mask 变成更像真实笔迹的 alpha。
    """
    profile = get_pen_profile(settings)

    pressure_mask = create_pressure_mask(
        text_mask,
        settings,
        rng,
    )

    texture_mask = create_stroke_texture_mask(
        pressure_mask.size,
        settings,
        rng,
    )

    alpha = ImageChops.multiply(
        pressure_mask,
        texture_mask,
    )

    if settings.pen_style == "钢笔" and float(profile["feather"]) > 0:
        feather = alpha.filter(
            ImageFilter.GaussianBlur(radius=1.0)
        ).point(
            lambda pixel: int(pixel * float(profile["feather"]))
        )
        alpha = ImageChops.lighter(alpha, feather)

    color = varied_ink_color(settings, rng)
    line_alpha = rng.randint(
        int(profile["alpha_low"]),
        int(profile["alpha_high"]),
    )

    return alpha, color, line_alpha


def cleanup_connection_hairlines(
    alpha_mask: Image.Image,
) -> Image.Image:
    """
    清理文字周围低透明度的小细线和毛刺。

    先删除非常淡的孤立像素，再轻微中值滤波。
    不会人工连接字符，也不会增加额外入笔线。
    """
    thresholded = alpha_mask.point(
        lambda pixel: 0 if pixel < 26 else pixel
    )

    smoothed = thresholded.filter(
        ImageFilter.MedianFilter(3)
    )

    cleaned = Image.blend(
        thresholded,
        smoothed,
        0.38,
    ).point(
        lambda pixel: 0 if pixel < 18 else pixel
    )

    return cleaned


def adjust_stroke_weight(
    alpha_mask: Image.Image,
    font_weight: float,
) -> Image.Image:
    """
    调整字体粗细。
    1.0 为默认，>1 更粗，<1 更细。
    """
    if abs(font_weight - 1.0) < 0.001:
        return alpha_mask

    result = alpha_mask

    if font_weight > 1.0:
        remaining = font_weight - 1.0

        while remaining > 0.001:
            expanded = result.filter(
                ImageFilter.MaxFilter(3)
            )
            blend_amount = min(0.82, remaining / 0.22)
            result = Image.blend(
                result,
                expanded,
                blend_amount,
            )
            remaining -= 0.22

    else:
        remaining = 1.0 - font_weight

        while remaining > 0.001:
            shrunk = result.filter(
                ImageFilter.MinFilter(3)
            )
            blend_amount = min(0.75, remaining / 0.25)
            result = Image.blend(
                result,
                shrunk,
                blend_amount,
            )
            remaining -= 0.25

    result = result.filter(
        ImageFilter.GaussianBlur(radius=0.25)
    ).point(
        lambda pixel: 0 if pixel < 8 else pixel
    )

    return result


# ============================================================
# 文档文字提取
# ============================================================
def extract_text_from_txt(file_bytes: bytes) -> str:
    encodings = [
        "utf-8-sig",
        "utf-8",
        "cp1251",
        "windows-1251",
    ]

    for encoding in encodings:
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise ValueError("无法识别 TXT 文件编码。请将文件另存为 UTF-8。")


def extract_text_from_docx(file_bytes: bytes) -> str:
    document = Document(BytesIO(file_bytes))
    result: list[str] = []

    for paragraph in document.paragraphs:
        result.append(paragraph.text)

    for table in document.tables:
        result.append("")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            result.append(" | ".join(cells))

    return "\n".join(result)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    result: list[str] = []

    with fitz.open(stream=file_bytes, filetype="pdf") as document:
        for page_index, page in enumerate(document):
            page_text = page.get_text("text").strip()

            if page_text:
                result.append(page_text)

            if page_index < len(document) - 1:
                result.append("")

    return "\n".join(result)


def extract_uploaded_file(uploaded_file) -> str:
    file_bytes = uploaded_file.getvalue()
    suffix = Path(uploaded_file.name).suffix.lower()

    if suffix == ".txt":
        return extract_text_from_txt(file_bytes)

    if suffix == ".docx":
        return extract_text_from_docx(file_bytes)

    if suffix == ".pdf":
        return extract_text_from_pdf(file_bytes)

    raise ValueError("暂不支持这个文件格式。")


# ============================================================
# 作业本纸张
# ============================================================
def add_light_paper_texture(
    image: Image.Image,
    settings: RenderSettings,
    rng: random.Random,
) -> None:
    """
    添加更明显但仍自然的纸张质感：
    - 细颗粒
    - 大面积明暗变化
    - 轻微纤维感
    - 淡旧化污渍
    """
    draw = ImageDraw.Draw(image)
    age = settings.paper_age

    grain_count = 1800 + int(age * 4200)

    for _ in range(grain_count):
        x = rng.randint(0, PAGE_WIDTH - 1)
        y = rng.randint(0, PAGE_HEIGHT - 1)
        spread = 2 + int(age * 7)
        shade = rng.randint(-spread, spread)

        draw.point(
            (x, y),
            fill=(
                clamp(248 + shade),
                clamp(246 + shade),
                clamp(239 + shade - int(age * 4)),
            ),
        )

    # 大面积纸色起伏，让纸更有“面”的感觉
    cloud_layer = Image.new("RGBA", image.size, (0, 0, 0, 0))
    cloud_draw = ImageDraw.Draw(cloud_layer)

    cloud_count = 8 + int(age * 12)
    for _ in range(cloud_count):
        cx = rng.randint(0, PAGE_WIDTH)
        cy = rng.randint(0, PAGE_HEIGHT)
        rx = rng.randint(70, 220)
        ry = rng.randint(50, 180)

        tint = rng.choice(
            [
                (255, 253, 247, rng.randint(6, 12)),
                (236, 231, 220, rng.randint(5, 11)),
                (244, 239, 229, rng.randint(4, 10)),
            ]
        )

        cloud_draw.ellipse(
            (cx - rx, cy - ry, cx + rx, cy + ry),
            fill=tint,
        )

    cloud_layer = cloud_layer.filter(
        ImageFilter.GaussianBlur(radius=36)
    )
    image.alpha_composite(cloud_layer)

    # 使用柔和颗粒模拟纸纤维，避免生成随机短细线。
    fiber_noise = Image.effect_noise(
        image.size,
        5.0 + age * 7.0,
    ).convert("L")

    fiber_noise = fiber_noise.filter(
        ImageFilter.GaussianBlur(radius=0.65)
    )

    fiber_alpha = fiber_noise.point(
        lambda pixel: int(
            abs(pixel - 128)
            * (0.10 + age * 0.10)
        )
    )

    fiber_layer = Image.new(
        "RGBA",
        image.size,
        (225, 220, 210, 0),
    )
    fiber_layer.putalpha(fiber_alpha)
    image.alpha_composite(fiber_layer)

    if age > 0.16:
        stain_layer = Image.new(
            "RGBA",
            image.size,
            (0, 0, 0, 0),
        )
        stain_draw = ImageDraw.Draw(stain_layer)

        stain_count = 1 + int((age - 0.15) * 5)

        for _ in range(stain_count):
            cx = rng.randint(120, PAGE_WIDTH - 120)
            cy = rng.randint(120, PAGE_HEIGHT - 120)
            rx = rng.randint(40, 110)
            ry = rng.randint(24, 78)

            stain_draw.ellipse(
                (cx - rx, cy - ry, cx + rx, cy + ry),
                fill=(158, 128, 88, rng.randint(2, 8)),
            )

        stain_layer = stain_layer.filter(
            ImageFilter.GaussianBlur(radius=26)
        )
        image.alpha_composite(stain_layer)


def apply_paper_wrinkles(
    image: Image.Image,
    settings: RenderSettings,
    rng: random.Random,
) -> None:
    """
    添加纸张褶皱：
    - 用淡阴影 + 高光模拟折痕
    - 强度较低时非常隐约
    - 适合拍照或旧化纸面
    """
    wrinkle_strength = settings.wrinkle_strength

    if wrinkle_strength <= 0.01:
        return

    wrinkle_layer = Image.new(
        "RGBA",
        image.size,
        (0, 0, 0, 0),
    )
    wrinkle_draw = ImageDraw.Draw(wrinkle_layer)

    wrinkle_count = max(
        1,
        int(1 + wrinkle_strength * 5),
    )

    for _ in range(wrinkle_count):
        orientation = rng.choices(
            ["vertical", "horizontal", "diagonal"],
            weights=[42, 18, 40],
            k=1,
        )[0]

        if orientation == "vertical":
            x = rng.randint(120, PAGE_WIDTH - 120)
            start = (x + rng.randint(-18, 18), 0)
            end = (x + rng.randint(-18, 18), PAGE_HEIGHT)
            control = (
                x + rng.randint(-70, 70),
                PAGE_HEIGHT // 2 + rng.randint(-220, 220),
            )
        elif orientation == "horizontal":
            y = rng.randint(180, PAGE_HEIGHT - 180)
            start = (0, y + rng.randint(-12, 12))
            end = (PAGE_WIDTH, y + rng.randint(-12, 12))
            control = (
                PAGE_WIDTH // 2 + rng.randint(-240, 240),
                y + rng.randint(-60, 60),
            )
        else:
            if rng.random() < 0.5:
                start = (rng.randint(0, 120), rng.randint(0, PAGE_HEIGHT // 3))
                end = (PAGE_WIDTH - rng.randint(0, 120), PAGE_HEIGHT - rng.randint(0, PAGE_HEIGHT // 3))
            else:
                start = (rng.randint(0, 120), PAGE_HEIGHT - rng.randint(0, PAGE_HEIGHT // 3))
                end = (PAGE_WIDTH - rng.randint(0, 120), rng.randint(0, PAGE_HEIGHT // 3))

            control = (
                PAGE_WIDTH // 2 + rng.randint(-220, 220),
                PAGE_HEIGHT // 2 + rng.randint(-220, 220),
            )

        points = sample_quadratic_curve(
            start,
            control,
            end,
            steps=80,
        )

        shadow_alpha = int(18 + wrinkle_strength * 38)
        highlight_alpha = int(14 + wrinkle_strength * 28)

        # 阴影
        shadow_offset = rng.randint(1, 3)
        shadow_points = [(px + shadow_offset, py + shadow_offset) for px, py in points]
        wrinkle_draw.line(
            shadow_points,
            fill=(150, 140, 130, shadow_alpha),
            width=1 + int(wrinkle_strength * 2),
            joint="curve",
        )

        # 高光
        highlight_offset = rng.randint(1, 2)
        highlight_points = [(px - highlight_offset, py - highlight_offset) for px, py in points]
        wrinkle_draw.line(
            highlight_points,
            fill=(255, 254, 249, highlight_alpha),
            width=1 + int(wrinkle_strength),
            joint="curve",
        )

        # 中间轻微折带
        if rng.random() < 0.65:
            wrinkle_draw.line(
                points,
                fill=(210, 202, 190, int(10 + wrinkle_strength * 16)),
                width=1,
                joint="curve",
            )

    wrinkle_layer = wrinkle_layer.filter(
        ImageFilter.GaussianBlur(radius=1.2 + wrinkle_strength * 1.4)
    )
    image.alpha_composite(wrinkle_layer)


def create_paper_background(
    settings: RenderSettings,
    rng: random.Random,
) -> Image.Image:
    """
    生成学生作业本纸张。
    """
    warmth = int(settings.paper_age * 12)

    image = Image.new(
        "RGBA",
        (PAGE_WIDTH, PAGE_HEIGHT),
        (
            249 - warmth,
            247 - warmth,
            241 - int(warmth * 0.7),
            255,
        ),
    )

    add_light_paper_texture(
        image,
        settings,
        rng,
    )

    apply_paper_wrinkles(
        image,
        settings,
        rng,
    )

    draw = ImageDraw.Draw(image)

    ruled_types = {
        "俄语作业本",
        "普通横线本",
    }

    if settings.paper_type in ruled_types:
        if settings.paper_type == "俄语作业本":
            # 顶部双线和左侧字段，接近俄语学生作业本。
            draw.line(
                (48, 62, PAGE_WIDTH - 48, 62),
                fill=(169, 180, 210, 220),
                width=2,
            )
            draw.line(
                (48, 78, PAGE_WIDTH - 48, 78),
                fill=(190, 198, 219, 190),
                width=1,
            )

            draw.text(
                (57, 31),
                "Дата",
                fill=(132, 143, 168, 205),
            )
            draw.line(
                (92, 48, 225, 48),
                fill=(155, 165, 192, 200),
                width=1,
            )

            draw.text(
                (PAGE_WIDTH - 285, 31),
                "Классная работа",
                fill=(132, 143, 168, 190),
            )

        for y in range(
            TOP_MARGIN,
            PAGE_HEIGHT - BOTTOM_MARGIN,
            settings.rule_spacing,
        ):
            line_color = (
                (164, 181, 216, 215)
                if settings.paper_type == "俄语作业本"
                else (177, 191, 220, 205)
            )

            draw.line(
                (50, y, PAGE_WIDTH - 50, y),
                fill=line_color,
                width=2,
            )

        red_line_x = LEFT_MARGIN + 48

        draw.line(
            (
                red_line_x,
                68,
                red_line_x,
                PAGE_HEIGHT - 68,
            ),
            fill=(202, 132, 137, 205),
            width=2,
        )

    elif settings.paper_type == "方格本":
        spacing = settings.rule_spacing

        for x in range(
            50,
            PAGE_WIDTH - 50,
            spacing,
        ):
            draw.line(
                (x, 50, x, PAGE_HEIGHT - 50),
                fill=(202, 213, 231, 175),
                width=1,
            )

        for y in range(
            50,
            PAGE_HEIGHT - 50,
            spacing,
        ):
            draw.line(
                (50, y, PAGE_WIDTH - 50, y),
                fill=(202, 213, 231, 175),
                width=1,
            )

    return image


def get_text_start_x(
    settings: RenderSettings,
) -> int:
    if settings.paper_type in {
        "俄语作业本",
        "普通横线本",
    }:
        return LEFT_MARGIN + 48 + 20

    return LEFT_MARGIN


def draw_notebook_header(
    page: Image.Image,
    settings: RenderSettings,
    rng: random.Random,
) -> None:
    """
    在页面顶部绘制日期和课题。
    """
    if not settings.header_enabled:
        return

    header_settings = replace(
        settings,
        font_size=max(30, settings.font_size - 10),
        texture_strength=min(
            1.0,
            settings.texture_strength + 0.04,
        ),
        randomness=min(
            2,
            settings.randomness,
        ),
        baseline_wave=0.25,
        flourish_level=0,
        correction_level=0,
        teacher_marks=0,
        header_enabled=False,
        show_page_number=False,
    )

    baseline_y = 92

    if settings.header_date.strip():
        draw_handwritten_line(
            page=page,
            text=settings.header_date.strip(),
            baseline_y=baseline_y,
            x_start=72,
            settings=header_settings,
            rng=rng,
        )

    if settings.header_lesson.strip():
        lesson_x = max(
            360,
            PAGE_WIDTH // 2 - 90,
        )

        draw_handwritten_line(
            page=page,
            text=settings.header_lesson.strip(),
            baseline_y=baseline_y,
            x_start=lesson_x,
            settings=header_settings,
            rng=rng,
        )


def draw_page_number(
    page: Image.Image,
    page_number: int,
    settings: RenderSettings,
    rng: random.Random,
) -> None:
    if not settings.show_page_number:
        return

    number_settings = replace(
        settings,
        font_size=28,
        randomness=1,
        baseline_wave=0.0,
        flourish_level=0,
        correction_level=0,
        teacher_marks=0,
        header_enabled=False,
        show_page_number=False,
    )

    draw_handwritten_line(
        page=page,
        text=str(page_number),
        baseline_y=PAGE_HEIGHT - 48,
        x_start=PAGE_WIDTH - 105,
        settings=number_settings,
        rng=rng,
    )


def apply_page_finish(
    image: Image.Image,
    settings: RenderSettings,
    rng: random.Random,
) -> Image.Image:
    """
    添加扫描件或手机拍照外观。
    """
    page = image.convert("RGB")

    if settings.photo_effect == "无":
        return page

    if settings.photo_effect == "扫描件":
        page = ImageOps.autocontrast(
            page,
            cutoff=0.4,
        )

        page = ImageEnhance.Contrast(
            page
        ).enhance(1.035)

        page = ImageEnhance.Sharpness(
            page
        ).enhance(1.08)

        vignette = Image.new(
            "L",
            page.size,
            255,
        )
        vignette_draw = ImageDraw.Draw(vignette)

        for inset in range(0, 65, 5):
            alpha = clamp(
                224 + inset // 2
            )

            vignette_draw.rectangle(
                (
                    inset,
                    inset,
                    PAGE_WIDTH - inset - 1,
                    PAGE_HEIGHT - inset - 1,
                ),
                outline=alpha,
                width=5,
            )

        shadow = Image.new(
            "RGB",
            page.size,
            (226, 226, 226),
        )

        return Image.composite(
            page,
            shadow,
            vignette,
        )

    # 手机拍照：桌面背景、页面阴影和轻微旋转。
    canvas_width = 1420
    canvas_height = 1940

    table = Image.new(
        "RGB",
        (canvas_width, canvas_height),
        (48, 47, 45),
    )

    table_draw = ImageDraw.Draw(table)

    for y in range(0, canvas_height, 7):
        shade = rng.randint(-5, 6)
        table_draw.line(
            (0, y, canvas_width, y),
            fill=(
                clamp(48 + shade),
                clamp(47 + shade),
                clamp(45 + shade),
            ),
            width=1,
        )

    angle = rng.uniform(-0.55, 0.55)

    page_rgba = page.convert("RGBA").rotate(
        angle,
        expand=True,
        resample=Image.Resampling.BICUBIC,
        fillcolor=(0, 0, 0, 0),
    )

    shadow_alpha = page_rgba.getchannel("A").filter(
        ImageFilter.GaussianBlur(radius=16)
    )

    shadow_layer = Image.new(
        "RGBA",
        page_rgba.size,
        (0, 0, 0, 0),
    )
    shadow_layer.putalpha(
        shadow_alpha.point(
            lambda pixel: int(pixel * 0.40)
        )
    )

    paste_x = (
        canvas_width - page_rgba.width
    ) // 2
    paste_y = (
        canvas_height - page_rgba.height
    ) // 2

    table_rgba = table.convert("RGBA")

    table_rgba.alpha_composite(
        shadow_layer,
        (paste_x + 15, paste_y + 18),
    )

    table_rgba.alpha_composite(
        page_rgba,
        (paste_x, paste_y),
    )

    return table_rgba.convert("RGB")


# ============================================================
# 自动换行和分页
# ============================================================
def text_width(
    text: str,
    font: ImageFont.FreeTypeFont,
) -> float:
    measuring_image = Image.new("RGB", (10, 10), "white")
    draw = ImageDraw.Draw(measuring_image)
    return draw.textlength(text, font=font)


def split_long_word(
    word: str,
    font: ImageFont.FreeTypeFont,
    max_width: int,
) -> list[str]:
    parts: list[str] = []
    current = ""

    for char in word:
        candidate = current + char
        if text_width(candidate, font) <= max_width:
            current = candidate
        else:
            if current:
                parts.append(current)
            current = char

    if current:
        parts.append(current)

    return parts


def wrap_paragraph(
    paragraph: str,
    font: ImageFont.FreeTypeFont,
    max_width: int,
) -> list[str]:
    if not paragraph.strip():
        return [""]

    words = paragraph.split()
    lines: list[str] = []
    current = ""

    for word in words:
        candidate = word if not current else f"{current} {word}"

        if text_width(candidate, font) <= max_width:
            current = candidate
            continue

        if current:
            lines.append(current)
            current = ""

        if text_width(word, font) <= max_width:
            current = word
            continue

        parts = split_long_word(word, font, max_width)
        if parts:
            lines.extend(parts[:-1])
            current = parts[-1]

    if current:
        lines.append(current)

    return lines


def wrap_text(
    text: str,
    font: ImageFont.FreeTypeFont,
    max_width: int,
) -> list[str]:
    all_lines: list[str] = []

    for paragraph in normalize_text(text).split("\n"):
        all_lines.extend(
            wrap_paragraph(
                paragraph=paragraph,
                font=font,
                max_width=max_width,
            )
        )

    return all_lines


def paginate_lines(
    lines: list[str],
    settings: RenderSettings,
) -> list[list[str]]:
    available_height = PAGE_HEIGHT - TOP_MARGIN - BOTTOM_MARGIN
    maximum_rows = max(1, available_height // settings.rule_spacing)

    pages: list[list[str]] = []
    current_page: list[str] = []

    for line in lines:
        if len(current_page) >= maximum_rows:
            pages.append(current_page)
            current_page = []

        current_page.append(line)

    if current_page or not pages:
        pages.append(current_page)

    return pages


# ============================================================
# 手写字渲染
# ============================================================
def shear_right(
    image: Image.Image,
    shear: float,
) -> Image.Image:
    if shear <= 0:
        return image

    extra_width = int(abs(shear) * image.height)
    new_width = image.width + extra_width
    offset = -shear * image.height

    return image.transform(
        (new_width, image.height),
        Image.Transform.AFFINE,
        (1, shear, offset, 0, 1, 0),
        resample=Image.Resampling.BICUBIC,
    )


def sample_quadratic_curve(
    start: tuple[float, float],
    control: tuple[float, float],
    end: tuple[float, float],
    steps: int = 18,
) -> list[tuple[int, int]]:
    """
    生成二次贝塞尔曲线上的点。
    """
    points: list[tuple[int, int]] = []

    for index in range(steps + 1):
        t = index / steps
        one_minus_t = 1.0 - t

        x = (
            one_minus_t * one_minus_t * start[0]
            + 2 * one_minus_t * t * control[0]
            + t * t * end[0]
        )

        y = (
            one_minus_t * one_minus_t * start[1]
            + 2 * one_minus_t * t * control[1]
            + t * t * end[1]
        )

        points.append((round(x), round(y)))

    return points


def draw_connector_curve(
    page: Image.Image,
    start_x: int,
    start_y: int,
    end_x: int,
    end_y: int,
    color: tuple[int, int, int],
    alpha: int,
    strength: float,
    rng: random.Random,
) -> None:
    """
    用弧线连接相邻单词，让整行更像连续书写。
    """
    gap = end_x - start_x

    if gap <= 0 or gap > 44:
        return

    lift = rng.uniform(0.5, 3.5) * max(0.2, strength)
    control_x = start_x + gap * rng.uniform(0.42, 0.58)
    control_y = min(start_y, end_y) - lift

    points = sample_quadratic_curve(
        (start_x, start_y),
        (control_x, control_y),
        (end_x, end_y),
        steps=max(10, gap),
    )

    draw = ImageDraw.Draw(page)

    draw.line(
        points,
        fill=(color[0], color[1], color[2], alpha),
        width=1,
        joint="curve",
    )

    if strength > 0.72 and rng.random() < 0.35:
        second_points = [
            (x, y + 1)
            for x, y in points[:-2]
        ]

        draw.line(
            second_points,
            fill=(
                color[0],
                color[1],
                color[2],
                max(55, alpha - 115),
            ),
            width=1,
            joint="curve",
        )


def draw_entry_stroke(
    page: Image.Image,
    x: int,
    baseline_y: int,
    color: tuple[int, int, int],
    strength: float,
    rng: random.Random,
) -> None:
    """
    在一行开头偶尔添加轻微入笔。
    """
    if strength < 0.45 or rng.random() > 0.42:
        return

    length = rng.randint(7, 15)
    start = (x - length, baseline_y + rng.randint(1, 4))
    control = (x - length // 2, baseline_y - rng.randint(0, 3))
    end = (x + 1, baseline_y)

    points = sample_quadratic_curve(
        start,
        control,
        end,
        steps=14,
    )

    ImageDraw.Draw(page).line(
        points,
        fill=(color[0], color[1], color[2], 175),
        width=1,
        joint="curve",
    )


def draw_end_flourish(
    page: Image.Image,
    start_x: int,
    baseline_y: int,
    color: tuple[int, int, int],
    level: int,
    line_text: str,
    rng: random.Random,
) -> None:
    """
    在句尾或行尾添加自然收笔弧线。
    """
    if level <= 0 or start_x >= PAGE_WIDTH - RIGHT_MARGIN - 8:
        return

    base_probability = {
        1: 0.20,
        2: 0.43,
        3: 0.70,
    }.get(level, 0.20)

    if line_text.rstrip().endswith((".", "!", "?", "…", ":", ";")):
        base_probability += 0.18

    if rng.random() > min(0.92, base_probability):
        return

    available = PAGE_WIDTH - RIGHT_MARGIN - start_x
    length = min(
        available,
        rng.randint(12 + level * 5, 23 + level * 11),
    )

    if length < 7:
        return

    end_x = start_x + length
    end_y = baseline_y + rng.randint(-3, 3)

    control_x = start_x + length * rng.uniform(0.38, 0.64)
    control_y = baseline_y + rng.randint(-7, 3)

    points = sample_quadratic_curve(
        (start_x, baseline_y),
        (control_x, control_y),
        (end_x, end_y),
        steps=22,
    )

    draw = ImageDraw.Draw(page)
    draw.line(
        points,
        fill=(color[0], color[1], color[2], rng.randint(155, 215)),
        width=1,
        joint="curve",
    )

    if level >= 3 and rng.random() < 0.32:
        loop_radius = rng.randint(4, 8)
        draw.arc(
            (
                end_x - loop_radius,
                end_y - loop_radius,
                end_x + loop_radius,
                end_y + loop_radius,
            ),
            start=rng.randint(145, 190),
            end=rng.randint(315, 355),
            fill=(color[0], color[1], color[2], 145),
            width=1,
        )


def get_correction_stroke_width(
    settings: RenderSettings,
) -> int:
    """
    让涂改笔画粗细与正文同步。

    字体越大、字体粗细越高，涂改笔画也越粗；
    不同书写工具只做很小的自然差异。
    """
    pen_factor = {
        "铅笔": 0.92,
        "圆珠笔": 0.98,
        "钢笔": 1.04,
        "中性笔": 1.08,
    }.get(settings.pen_style, 1.0)

    width = round(
        settings.font_size
        / 22.0
        * settings.font_weight
        * pen_factor
    )

    return max(1, min(6, width))


def correction_fill(
    color: tuple[int, int, int],
    rng: random.Random,
) -> tuple[int, int, int, int]:
    """
    涂改使用接近正文的不透明墨迹，不再淡化。
    """
    return (
        color[0],
        color[1],
        color[2],
        rng.randint(238, 255),
    )


def draw_strikethrough(
    page: Image.Image,
    left: int,
    top: int,
    width: int,
    height: int,
    color: tuple[int, int, int],
    settings: RenderSettings,
    rng: random.Random,
    start_ratio: float = 0.0,
    end_ratio: float = 1.0,
) -> None:
    """
    绘制一到两道自然删除线。
    """
    draw = ImageDraw.Draw(page)
    stroke_width = get_correction_stroke_width(settings)

    start_ratio = max(0.0, min(1.0, start_ratio))
    end_ratio = max(start_ratio + 0.05, min(1.0, end_ratio))

    line_left = left + int(width * start_ratio)
    line_right = left + int(width * end_ratio)
    line_width = max(8, line_right - line_left)

    line_count = rng.choices(
        [1, 2],
        weights=[72, 28],
        k=1,
    )[0]

    for index in range(line_count):
        start_y = (
            top
            + int(height * rng.uniform(0.40, 0.64))
            + index * rng.randint(1, max(2, stroke_width))
        )
        end_y = start_y + rng.randint(-4, 4)

        middle_x = line_left + line_width // 2
        middle_y = (
            (start_y + end_y) // 2
            + rng.randint(-3, 3)
        )

        points = sample_quadratic_curve(
            (line_left - 2, start_y),
            (middle_x, middle_y),
            (line_right + 2, end_y),
            steps=max(10, line_width // 5),
        )

        draw.line(
            points,
            fill=correction_fill(color, rng),
            width=stroke_width,
            joint="curve",
        )


def choose_error_fragment(
    word: str,
    left: int,
    width: int,
    rng: random.Random,
) -> dict[str, object]:
    length = len(word)

    if length <= 3:
        start_index = 0
        end_index = length
    else:
        region = rng.choices(
            ["prefix", "middle", "suffix"],
            weights=[12, 26, 62],
            k=1,
        )[0]

        frag_len = min(
            length,
            rng.randint(1, min(3, length)),
        )

        if region == "prefix":
            start_index = 0
            end_index = frag_len
        elif region == "middle":
            max_start = max(1, length - frag_len - 1)
            start_index = rng.randint(1, max_start)
            end_index = start_index + frag_len
        else:
            start_index = max(0, length - frag_len)
            end_index = length

    start_ratio = start_index / max(1, length)
    end_ratio = end_index / max(1, length)

    fragment_left = left + int(width * start_ratio)
    fragment_width = max(
        10,
        int(width * (end_ratio - start_ratio)),
    )

    return {
        "text": word[start_index:end_index],
        "start_ratio": start_ratio,
        "end_ratio": end_ratio,
        "left": fragment_left,
        "width": fragment_width,
        "start_index": start_index,
        "end_index": end_index,
    }


def draw_scribble_correction(
    page: Image.Image,
    left: int,
    top: int,
    width: int,
    height: int,
    color: tuple[int, int, int],
    settings: RenderSettings,
    rng: random.Random,
    style: str = "messy",
) -> None:
    """
    规整斜线或密集乱涂。
    所有线条都与正文字体粗细同步。
    """
    draw = ImageDraw.Draw(page)
    stroke_width = get_correction_stroke_width(settings)

    if style == "neat":
        primary_count = rng.randint(2, 4)

        for index in range(primary_count):
            x1 = left + rng.randint(-1, max(2, width // 5))
            y1 = (
                top
                + int(height * rng.uniform(0.24, 0.38))
                + index * rng.randint(1, max(2, stroke_width))
            )
            x2 = left + width - rng.randint(0, max(2, width // 7))
            y2 = (
                top
                + int(height * rng.uniform(0.66, 0.82))
                + index * rng.randint(-1, 2)
            )

            draw.line(
                (x1, y1, x2, y2),
                fill=correction_fill(color, rng),
                width=stroke_width,
                joint="curve",
            )

        if rng.random() < 0.58:
            for _ in range(rng.randint(1, 2)):
                x1 = left + rng.randint(-1, max(2, width // 6))
                y1 = top + int(height * rng.uniform(0.67, 0.82))
                x2 = left + width - rng.randint(0, max(2, width // 7))
                y2 = top + int(height * rng.uniform(0.24, 0.41))

                draw.line(
                    (x1, y1, x2, y2),
                    fill=correction_fill(color, rng),
                    width=stroke_width,
                    joint="curve",
                )
        return

    stroke_count = rng.randint(5, 9)

    for _ in range(stroke_count):
        points: list[tuple[int, int]] = []
        segments = rng.randint(3, 6)

        if rng.random() < 0.5:
            x = left + rng.randint(-1, max(1, width // 7))
        else:
            x = left + width - rng.randint(0, max(1, width // 7))

        y = top + int(height * rng.uniform(0.22, 0.78))
        points.append((x, y))

        for _ in range(segments):
            x = left + rng.randint(0, max(2, width))
            y = top + rng.randint(
                int(height * 0.18),
                max(int(height * 0.82), 1),
            )
            points.append((x, y))

        draw.line(
            points,
            fill=correction_fill(color, rng),
            width=stroke_width,
            joint="curve",
        )


def draw_caret_mark(
    page: Image.Image,
    center_x: int,
    baseline_y: int,
    color: tuple[int, int, int],
    settings: RenderSettings,
    rng: random.Random,
) -> None:
    stroke_width = get_correction_stroke_width(settings)
    size = rng.randint(6, 10) + stroke_width
    bottom_y = baseline_y + rng.randint(7, 11)
    top_y = bottom_y - size

    ImageDraw.Draw(page).line(
        (
            center_x - size,
            bottom_y,
            center_x,
            top_y,
            center_x + size,
            bottom_y,
        ),
        fill=correction_fill(color, rng),
        width=stroke_width,
        joint="curve",
    )


def draw_rewrite_above(
    page: Image.Image,
    rewrite_text: str,
    left: int,
    top: int,
    settings: RenderSettings,
    rng: random.Random,
    compact: bool = False,
) -> None:
    """
    上方重写保留当前字体粗细和书写工具，不再降低透明度。
    """
    rewrite_settings = replace(
        settings,
        font_size=max(
            22 if compact else 24,
            settings.font_size - (10 if compact else 8),
        ),
        texture_strength=settings.texture_strength,
        font_weight=settings.font_weight,
        randomness=max(0, settings.randomness - 1),
        correction_level=0,
        teacher_marks=0,
        flourish_level=0,
        baseline_wave=0.0,
    )

    rewrite_image, _, _, _ = render_word_image(
        rewrite_text,
        rewrite_settings,
        rng,
    )

    rewrite_x = left + rng.randint(-2, 8)
    rewrite_y = (
        top
        - int(
            settings.font_size
            * (0.42 if compact else 0.48)
        )
        + rng.randint(-4, 3)
    )

    page.alpha_composite(
        rewrite_image,
        (
            max(0, rewrite_x),
            max(0, rewrite_y),
        ),
    )


def draw_inline_overwrite(
    page: Image.Image,
    text: str,
    left: int,
    top: int,
    settings: RenderSettings,
    rng: random.Random,
) -> None:
    """
    字上叠字使用正常墨迹，不再使用淡灰透明文字。
    """
    overwrite_settings = replace(
        settings,
        correction_level=0,
        teacher_marks=0,
        flourish_level=0,
        baseline_wave=0.0,
        randomness=max(0, settings.randomness - 1),
    )

    first, _, _, _ = render_word_image(
        text,
        overwrite_settings,
        rng,
    )

    first_alpha = first.getchannel("A").point(
        lambda pixel: int(
            pixel * rng.uniform(0.88, 1.0)
        )
    )
    first.putalpha(first_alpha)

    page.alpha_composite(
        first,
        (
            max(0, left + rng.randint(-2, 2)),
            max(0, top + rng.randint(-2, 1)),
        ),
    )

    if rng.random() < 0.42:
        second, _, _, _ = render_word_image(
            text,
            overwrite_settings,
            rng,
        )

        second_alpha = second.getchannel("A").point(
            lambda pixel: int(
                pixel * rng.uniform(0.82, 0.96)
            )
        )
        second.putalpha(second_alpha)

        page.alpha_composite(
            second,
            (
                max(0, left + rng.randint(-1, 3)),
                max(0, top + rng.randint(-1, 2)),
            ),
        )


def draw_diagonal_cancel(
    page: Image.Image,
    left: int,
    top: int,
    width: int,
    height: int,
    color: tuple[int, int, int],
    settings: RenderSettings,
    rng: random.Random,
) -> None:
    draw = ImageDraw.Draw(page)
    stroke_width = get_correction_stroke_width(settings)
    line_count = rng.choices(
        [1, 2],
        weights=[72, 28],
        k=1,
    )[0]

    x1 = left + rng.randint(-1, 2)
    y1 = top + int(height * rng.uniform(0.25, 0.45))
    x2 = left + width + rng.randint(-2, 2)
    y2 = top + int(height * rng.uniform(0.65, 0.82))

    draw.line(
        (x1, y1, x2, y2),
        fill=correction_fill(color, rng),
        width=stroke_width,
        joint="curve",
    )

    if line_count == 2:
        x3 = left + rng.randint(-1, 2)
        y3 = top + int(height * rng.uniform(0.67, 0.83))
        x4 = left + width + rng.randint(-2, 2)
        y4 = top + int(height * rng.uniform(0.22, 0.42))

        draw.line(
            (x3, y3, x4, y4),
            fill=correction_fill(color, rng),
            width=stroke_width,
            joint="curve",
        )


def draw_loop_cancel(
    page: Image.Image,
    left: int,
    top: int,
    width: int,
    height: int,
    color: tuple[int, int, int],
    settings: RenderSettings,
    rng: random.Random,
) -> None:
    draw = ImageDraw.Draw(page)
    stroke_width = get_correction_stroke_width(settings)

    cx = left + width // 2
    cy = top + int(height * 0.55)
    rx = max(8, width // 2 + rng.randint(0, 5))
    ry = max(6, int(height * 0.30) + rng.randint(0, 4))

    bbox = (
        cx - rx,
        cy - ry,
        cx + rx,
        cy + ry,
    )

    start_angle = rng.randint(150, 210)
    end_angle = start_angle + rng.randint(240, 325)

    draw.arc(
        bbox,
        start=start_angle,
        end=end_angle,
        fill=correction_fill(color, rng),
        width=stroke_width,
    )

    if rng.random() < 0.28:
        draw.arc(
            (
                bbox[0] + rng.randint(-2, 1),
                bbox[1] + rng.randint(-1, 2),
                bbox[2] + rng.randint(-1, 2),
                bbox[3] + rng.randint(-2, 1),
            ),
            start=start_angle + rng.randint(-18, 14),
            end=end_angle - rng.randint(12, 24),
            fill=correction_fill(color, rng),
            width=stroke_width,
        )


def draw_short_scratch(
    page: Image.Image,
    left: int,
    top: int,
    width: int,
    height: int,
    color: tuple[int, int, int],
    settings: RenderSettings,
    rng: random.Random,
) -> None:
    draw = ImageDraw.Draw(page)
    stroke_width = get_correction_stroke_width(settings)

    for _ in range(rng.randint(2, 4)):
        sx = left + rng.randint(-1, max(1, width // 4))
        ex = left + width - rng.randint(0, max(1, width // 5))
        sy = top + int(height * rng.uniform(0.38, 0.68))
        ey = sy + rng.randint(-3, 3)

        control_x = (sx + ex) // 2 + rng.randint(-4, 4)
        control_y = sy + rng.randint(-5, 5)

        points = sample_quadratic_curve(
            (sx, sy),
            (control_x, control_y),
            (ex, ey),
            steps=max(8, width // 5),
        )

        draw.line(
            points,
            fill=correction_fill(color, rng),
            width=stroke_width,
            joint="curve",
        )


def draw_fragment_hatch_fix(
    page: Image.Image,
    fragment_text: str,
    left: int,
    top: int,
    width: int,
    height: int,
    settings: RenderSettings,
    rng: random.Random,
    hatch_style: str,
) -> None:
    ink = varied_ink_color(settings, rng)

    if rng.random() < 0.58:
        draw_inline_overwrite(
            page,
            fragment_text,
            left,
            top,
            settings,
            rng,
        )

    style = hatch_style

    if style not in {
        "规整型",
        "乱涂型",
    }:
        style = rng.choices(
            ["规整型", "乱涂型", "短刮擦", "斜划"],
            weights=[35, 20, 28, 17],
            k=1,
        )[0]

    if style == "规整型":
        draw_scribble_correction(
            page,
            left,
            top,
            width,
            height,
            ink,
            settings,
            rng,
            style="neat",
        )
    elif style == "乱涂型":
        draw_scribble_correction(
            page,
            left,
            top,
            width,
            height,
            ink,
            settings,
            rng,
            style="messy",
        )
    elif style == "短刮擦":
        draw_short_scratch(
            page,
            left,
            top,
            width,
            height,
            ink,
            settings,
            rng,
        )
    else:
        draw_diagonal_cancel(
            page,
            left,
            top,
            width,
            height,
            ink,
            settings,
            rng,
        )

    if rng.random() < 0.86:
        draw_rewrite_above(
            page,
            fragment_text,
            left + rng.randint(-1, 3),
            top,
            settings,
            rng,
            compact=True,
        )


def draw_retry_block(
    page: Image.Image,
    text: str,
    left: int,
    top: int,
    settings: RenderSettings,
    rng: random.Random,
) -> None:
    ink = varied_ink_color(settings, rng)

    draw_scribble_correction(
        page,
        left,
        top,
        max(
            18,
            int(
                len(text)
                * settings.font_size
                * 0.45
            ),
        ),
        max(
            16,
            int(settings.font_size * 1.0),
        ),
        ink,
        settings,
        rng,
        style=(
            "messy"
            if rng.random() < 0.55
            else "neat"
        ),
    )

    if rng.random() < 0.82:
        draw_rewrite_above(
            page,
            text,
            left + rng.randint(-1, 4),
            top,
            settings,
            rng,
            compact=bool(rng.random() < 0.6),
        )


def choose_micro_fragment(
    word: str,
    left: int,
    width: int,
    rng: random.Random,
) -> dict[str, object]:
    length = len(word)

    if length <= 2:
        start_index = 0
        end_index = length
    else:
        region = rng.choices(
            ["suffix", "middle", "prefix"],
            weights=[62, 26, 12],
            k=1,
        )[0]

        max_fragment = 2 if length <= 5 else 3
        fragment_length = min(
            length,
            rng.randint(1, max_fragment),
        )

        if region == "suffix":
            start_index = max(
                0,
                length - fragment_length,
            )
            end_index = length
        elif region == "middle":
            max_start = max(
                1,
                length - fragment_length - 1,
            )
            start_index = rng.randint(
                1,
                max_start,
            )
            end_index = (
                start_index
                + fragment_length
            )
        else:
            start_index = 0
            end_index = fragment_length

    start_ratio = start_index / max(1, length)
    end_ratio = end_index / max(1, length)

    return {
        "text": word[start_index:end_index],
        "start_ratio": start_ratio,
        "end_ratio": end_ratio,
        "left": (
            left
            + int(width * start_ratio)
        ),
        "width": max(
            10,
            int(
                width
                * (end_ratio - start_ratio)
            ),
        ),
        "start_index": start_index,
        "end_index": end_index,
    }


def get_correction_style_weights(
    correction_style: str,
) -> dict[str, int]:
    if correction_style == "局部改错型":
        return {
            "tail_fix": 36,
            "inline_overwrite_fix": 22,
            "messy_hatch_fix": 7,
            "neat_hatch_fix": 18,
            "caret_insert": 8,
            "whole_word_retry": 3,
            "cross_fix": 6,
        }

    if correction_style == "规整型":
        return {
            "tail_fix": 22,
            "inline_overwrite_fix": 18,
            "messy_hatch_fix": 4,
            "neat_hatch_fix": 34,
            "caret_insert": 8,
            "whole_word_retry": 5,
            "cross_fix": 9,
        }

    if correction_style == "乱涂型":
        return {
            "tail_fix": 13,
            "inline_overwrite_fix": 12,
            "messy_hatch_fix": 35,
            "neat_hatch_fix": 5,
            "caret_insert": 6,
            "whole_word_retry": 17,
            "cross_fix": 12,
        }

    return {
        "tail_fix": 29,
        "inline_overwrite_fix": 18,
        "messy_hatch_fix": 14,
        "neat_hatch_fix": 13,
        "caret_insert": 10,
        "whole_word_retry": 8,
        "cross_fix": 8,
    }


def maybe_draw_correction(
    page: Image.Image,
    word: str,
    left: int,
    top: int,
    width: int,
    height: int,
    baseline_y: int,
    settings: RenderSettings,
    rng: random.Random,
) -> None:
    if settings.correction_level <= 0:
        return

    clean_word = word.strip(
        ".,!?;:…—-()[]{}«»\"'"
    )

    if len(clean_word) < 3:
        return

    base_denominator = {
        1: 36,
        2: 22,
        3: 13,
    }.get(settings.correction_level, 22)

    local_denominator = max(
        7,
        base_denominator
        + rng.randint(-3, 4)
        - (
            1
            if len(clean_word) >= 7
            else 0
        ),
    )

    if rng.randint(
        1,
        local_denominator,
    ) != 1:
        return

    color = varied_ink_color(
        settings,
        rng,
    )

    fragment = choose_micro_fragment(
        clean_word,
        left,
        width,
        rng,
    )

    fragment_left = int(fragment["left"])
    fragment_width = max(
        12,
        int(fragment["width"]),
    )
    fragment_text = (
        str(fragment["text"])
        or clean_word
    )

    weights = get_correction_style_weights(
        settings.correction_style
    )

    style = rng.choices(
        list(weights.keys()),
        weights=list(weights.values()),
        k=1,
    )[0]

    if style == "tail_fix":
        draw_fragment_hatch_fix(
            page,
            fragment_text,
            fragment_left,
            top,
            fragment_width,
            height,
            settings,
            rng,
            settings.correction_style,
        )

    elif style == "inline_overwrite_fix":
        draw_inline_overwrite(
            page,
            fragment_text,
            fragment_left,
            top,
            settings,
            rng,
        )

        if rng.random() < 0.80:
            draw_rewrite_above(
                page,
                fragment_text,
                fragment_left
                + rng.randint(-1, 2),
                top,
                settings,
                rng,
                compact=True,
            )

    elif style == "messy_hatch_fix":
        draw_scribble_correction(
            page,
            fragment_left,
            top,
            fragment_width,
            height,
            color,
            settings,
            rng,
            style="messy",
        )

        if rng.random() < 0.74:
            draw_rewrite_above(
                page,
                fragment_text,
                fragment_left
                + rng.randint(-1, 3),
                top,
                settings,
                rng,
                compact=True,
            )

    elif style == "neat_hatch_fix":
        draw_scribble_correction(
            page,
            fragment_left,
            top,
            fragment_width,
            height,
            color,
            settings,
            rng,
            style="neat",
        )

        if rng.random() < 0.86:
            draw_rewrite_above(
                page,
                fragment_text,
                fragment_left,
                top,
                settings,
                rng,
                compact=True,
            )

    elif style == "caret_insert":
        caret_x = (
            fragment_left
            + fragment_width // 2
        )

        draw_caret_mark(
            page,
            caret_x,
            baseline_y,
            color,
            settings,
            rng,
        )

        draw_rewrite_above(
            page,
            fragment_text,
            caret_x - rng.randint(5, 10),
            top,
            settings,
            rng,
            compact=True,
        )

    elif style == "whole_word_retry":
        draw_retry_block(
            page,
            clean_word,
            left,
            top,
            settings,
            rng,
        )

    else:
        if rng.random() < 0.58:
            draw_diagonal_cancel(
                page,
                fragment_left,
                top,
                fragment_width,
                height,
                color,
                settings,
                rng,
            )
        else:
            draw_short_scratch(
                page,
                fragment_left,
                top,
                fragment_width,
                height,
                color,
                settings,
                rng,
            )

        if rng.random() < 0.78:
            draw_rewrite_above(
                page,
                fragment_text,
                fragment_left
                + rng.randint(-1, 3),
                top,
                settings,
                rng,
                compact=True,
            )


def render_word_image(
    word: str,
    settings: RenderSettings,
    rng: random.Random,
) -> tuple[
    Image.Image,
    int,
    tuple[int, int, int],
    int,
]:
    """
    用当前笔迹工具渲染单个词，主要用于涂改后的上方重写。
    """
    scale = settings.quality_scale

    high_font = load_font(
        settings.font_path,
        settings.font_size * scale,
    )

    ascent, descent = high_font.getmetrics()
    padding = 12 * scale

    measuring_image = Image.new(
        "L",
        (1, 1),
        0,
    )
    measuring_draw = ImageDraw.Draw(
        measuring_image
    )

    try:
        bbox = measuring_draw.textbbox(
            (0, 0),
            word,
            font=high_font,
            anchor="ls",
            features=["calt", "liga", "clig", "kern"],
        )
    except Exception:
        bbox = measuring_draw.textbbox(
            (0, 0),
            word,
            font=high_font,
            anchor="ls",
        )

    content_width = max(
        1,
        bbox[2] - bbox[0],
    )

    canvas_width = (
        content_width
        + padding * 2
        + int(
            settings.slant
            * (ascent + descent)
        )
    )

    canvas_height = (
        ascent
        + descent
        + padding * 2
    )

    baseline_high = padding + ascent

    mask = Image.new(
        "L",
        (
            max(1, canvas_width),
            max(1, canvas_height),
        ),
        0,
    )

    draw = ImageDraw.Draw(mask)

    draw_text_with_cursive_features(
        draw,
        (
            padding - bbox[0],
            baseline_high,
        ),
        word,
        high_font,
        255,
        anchor="ls",
    )

    profile = get_pen_profile(settings)

    mask = mask.filter(
        ImageFilter.GaussianBlur(
            radius=float(
                profile["blur"]
            )
            * scale
        )
    )

    (
        alpha_mask,
        ink_color,
        word_alpha,
    ) = apply_pen_texture(
        mask,
        settings,
        rng,
    )

    alpha_mask = cleanup_connection_hairlines(
        alpha_mask
    )
    alpha_mask = adjust_stroke_weight(
        alpha_mask,
        settings.font_weight,
    )

    word_image = Image.new(
        "RGBA",
        mask.size,
        (
            ink_color[0],
            ink_color[1],
            ink_color[2],
            0,
        ),
    )

    word_image.putalpha(
        alpha_mask
    )

    width_change = rng.uniform(
        0.996
        - settings.randomness * 0.0008,
        1.004
        + settings.randomness * 0.0008,
    )

    height_change = rng.uniform(
        0.998
        - settings.randomness * 0.0005,
        1.002
        + settings.randomness * 0.0005,
    )

    target_width = max(
        1,
        int(
            word_image.width
            / scale
            * width_change
        ),
    )

    target_height = max(
        1,
        int(
            word_image.height
            / scale
            * height_change
        ),
    )

    baseline_after_resize = int(
        baseline_high
        / scale
        * height_change
    )

    word_image = word_image.resize(
        (
            target_width,
            target_height,
        ),
        Image.Resampling.LANCZOS,
    )

    word_image = shear_right(
        word_image,
        max(
            0.0,
            settings.slant
            + rng.uniform(-0.008, 0.008),
        ),
    )

    final_alpha = word_image.getchannel(
        "A"
    ).point(
        lambda pixel: int(
            pixel
            * word_alpha
            / 255
        )
    )

    word_image.putalpha(
        final_alpha
    )

    return (
        word_image,
        baseline_after_resize,
        ink_color,
        word_alpha,
    )


def warp_cursive_line(
    image: Image.Image,
    amplitude: float,
    rng: random.Random,
) -> Image.Image:
    """
    对整行文字进行非常轻微的纵向曲线变形。

    与逐词上下移动不同，这种方式不会切断单词内部的连笔，
    更接近参考图片中连续、平稳的俄语草写。
    """
    if amplitude <= 0.05 or image.width < 20:
        return image

    width, height = image.size
    segment_count = max(
        8,
        min(36, width // 45),
    )

    phase = rng.uniform(0.0, math.tau)
    frequency = rng.uniform(0.75, 1.25)
    drift = rng.uniform(-0.18, 0.18)

    boundaries = [
        round(index * width / segment_count)
        for index in range(segment_count + 1)
    ]

    offsets: list[float] = []

    for index in range(segment_count + 1):
        progress = index / segment_count

        sine_offset = (
            math.sin(
                phase
                + progress
                * math.tau
                * frequency
            )
            * amplitude
        )

        drift_offset = (
            drift
            * amplitude
            * (progress - 0.5)
        )

        offsets.append(
            sine_offset + drift_offset
        )

    mesh = []

    for index in range(segment_count):
        x0 = boundaries[index]
        x1 = boundaries[index + 1]

        if x1 <= x0:
            continue

        left_offset = offsets[index]
        right_offset = offsets[index + 1]

        mesh.append(
            (
                (x0, 0, x1, height),
                (
                    x0,
                    -left_offset,
                    x0,
                    height - left_offset,
                    x1,
                    height - right_offset,
                    x1,
                    -right_offset,
                ),
            )
        )

    return image.transform(
        image.size,
        Image.Transform.MESH,
        mesh,
        resample=Image.Resampling.BICUBIC,
    )


def render_cursive_line_image(
    text: str,
    settings: RenderSettings,
    rng: random.Random,
) -> tuple[
    Image.Image,
    int,
    tuple[int, int, int],
    int,
]:
    """
    把整行文字一次性绘制。

    这一版重点：
    1. 去掉字符间人工桥接，避免出现奇怪虚线；
    2. 保留字体自己的连字与连写；
    3. 增加钢笔 / 圆珠笔 / 中性笔 / 铅笔的纹理差异。
    """
    scale = settings.quality_scale

    high_font = load_font(
        settings.font_path,
        settings.font_size * scale,
    )

    ascent, descent = high_font.getmetrics()
    padding = 20 * scale
    baseline_high = padding + ascent

    measuring_image = Image.new(
        "L",
        (1, 1),
        0,
    )
    measuring_draw = ImageDraw.Draw(measuring_image)

    try:
        bbox = measuring_draw.textbbox(
            (0, 0),
            text,
            font=high_font,
            anchor="ls",
            features=["calt", "liga", "clig", "kern"],
        )
    except Exception:
        bbox = measuring_draw.textbbox(
            (0, 0),
            text,
            font=high_font,
            anchor="ls",
        )

    content_width = max(1, bbox[2] - bbox[0])

    canvas_width = (
        content_width
        + padding * 2
        + int(settings.slant * (ascent + descent))
    )

    canvas_height = ascent + descent + padding * 2

    mask = Image.new(
        "L",
        (max(1, canvas_width), max(1, canvas_height)),
        0,
    )

    mask_draw = ImageDraw.Draw(mask)

    draw_text_with_cursive_features(
        mask_draw,
        (padding - bbox[0], baseline_high),
        text,
        high_font,
        255,
        anchor="ls",
    )

    profile = get_pen_profile(settings)

    # 只做极轻微柔化，不再使用字符间桥接，避免虚线或怪异连接。
    blur_radius = float(profile["blur"]) * scale
    mask = mask.filter(
        ImageFilter.GaussianBlur(radius=blur_radius)
    )

    alpha_mask, ink_color, line_alpha = apply_pen_texture(
        mask,
        settings,
        rng,
    )

    alpha_mask = cleanup_connection_hairlines(
        alpha_mask
    )
    alpha_mask = adjust_stroke_weight(
        alpha_mask,
        settings.font_weight,
    )

    line_image = Image.new(
        "RGBA",
        mask.size,
        (ink_color[0], ink_color[1], ink_color[2], 0),
    )
    line_image.putalpha(alpha_mask)

    space_count = text.count(" ")

    width_factor = (
        1.0
        - settings.connection_strength * 0.012
        + rng.uniform(-0.002, 0.002)
    )

    target_width = int(
        line_image.width / scale * width_factor
        + settings.word_spacing * space_count
    )

    height_factor = rng.uniform(
        0.998 - settings.randomness * 0.0004,
        1.002 + settings.randomness * 0.0004,
    )

    target_height = int(
        line_image.height / scale * height_factor
    )

    target_width = max(1, target_width)
    target_height = max(1, target_height)

    baseline_after_resize = int(
        baseline_high / scale * height_factor
    )

    line_image = line_image.resize(
        (target_width, target_height),
        Image.Resampling.LANCZOS,
    )

    effective_slant = max(
        0.0,
        settings.slant + settings.connection_strength * 0.008 + rng.uniform(-0.005, 0.005),
    )
    line_image = shear_right(line_image, effective_slant)

    line_image = warp_cursive_line(
        line_image,
        amplitude=(settings.baseline_wave * 0.55),
        rng=rng,
    )

    rotation_angle = rng.uniform(
        -0.030 - settings.randomness * 0.015,
        0.030 + settings.randomness * 0.015,
    )
    old_height = line_image.height

    line_image = line_image.rotate(
        rotation_angle,
        expand=True,
        resample=Image.Resampling.BICUBIC,
    )

    baseline_after_resize += (
        line_image.height - old_height
    ) // 2

    # 总体透明度微调
    alpha_final = line_image.getchannel("A").point(
        lambda pixel: min(255, max(0, int(pixel * line_alpha / 255)))
    )
    line_image.putalpha(alpha_final)

    visible_bbox = line_image.getchannel("A").getbbox()
    if visible_bbox is not None:
        baseline_after_resize -= visible_bbox[1]
        line_image = line_image.crop(visible_bbox)

    return line_image, baseline_after_resize, ink_color, line_alpha


def approximate_word_positions(
    text: str,
    x_start: int,
    rendered_width: int,
    font: ImageFont.FreeTypeFont,
) -> list[tuple[str, int, int]]:
    """
    估算整行渲染后各单词的位置，用于添加涂改痕迹。

    这里只影响涂改位置，不影响主体文字连笔。
    """
    raw_width = max(
        1.0,
        text_width(text, font),
    )

    width_scale = (
        rendered_width / raw_width
    )

    positions: list[
        tuple[str, int, int]
    ] = []

    search_from = 0

    for word in text.split():
        word_index = text.find(
            word,
            search_from,
        )

        if word_index < 0:
            continue

        prefix = text[:word_index]

        word_left = (
            x_start
            + round(
                text_width(
                    prefix,
                    font,
                )
                * width_scale
            )
        )

        word_width = max(
            8,
            round(
                text_width(
                    word,
                    font,
                )
                * width_scale
            ),
        )

        positions.append(
            (
                word,
                word_left,
                word_width,
            )
        )

        search_from = (
            word_index + len(word)
        )

    return positions


def draw_handwritten_line(
    page: Image.Image,
    text: str,
    baseline_y: int,
    x_start: int,
    settings: RenderSettings,
    rng: random.Random,
) -> int:
    """
    整行一次性绘制，主要模拟参考图中的俄语连笔。

    单词内部由字体真正连接；
    单词之间保留自然空格，不再额外画生硬的连接线。
    """
    if not text.strip():
        return x_start

    (
        line_image,
        line_baseline,
        ink_color,
        _,
    ) = render_cursive_line_image(
        text=text,
        settings=settings,
        rng=rng,
    )

    paste_y = (
        baseline_y
        - line_baseline
    )

    max_allowed_width = (
        PAGE_WIDTH
        - RIGHT_MARGIN
        - x_start
    )

    if (
        line_image.width
        > max_allowed_width
        and max_allowed_width > 10
    ):
        resize_ratio = (
            max_allowed_width
            / line_image.width
        )

        old_height = (
            line_image.height
        )

        line_image = line_image.resize(
            (
                max_allowed_width,
                max(
                    1,
                    round(
                        old_height
                        * resize_ratio
                    ),
                ),
            ),
            Image.Resampling.LANCZOS,
        )

        line_baseline = round(
            line_baseline
            * resize_ratio
        )

        paste_y = (
            baseline_y
            - line_baseline
        )

    page.alpha_composite(
        line_image,
        (
            max(0, x_start),
            max(0, paste_y),
        ),
    )

    # 主体整行已经完成，再按估算位置添加少量涂改。
    normal_font = load_font(
        settings.font_path,
        settings.font_size,
    )

    word_positions = (
        approximate_word_positions(
            text=text,
            x_start=x_start,
            rendered_width=line_image.width,
            font=normal_font,
        )
    )

    for (
        word,
        word_left,
        word_width,
    ) in word_positions:
        maybe_draw_correction(
            page=page,
            word=word,
            left=word_left,
            top=max(0, paste_y),
            width=word_width,
            height=line_image.height,
            baseline_y=baseline_y,
            settings=settings,
            rng=rng,
        )

    line_end_x = (
        x_start + line_image.width
    )

    return line_end_x


def draw_teacher_feedback(
    page: Image.Image,
    line_baselines: list[int],
    text_start_x: int,
    settings: RenderSettings,
    rng: random.Random,
) -> None:
    """
    可选的老师红笔批改痕迹：
    对号、短下划线和圆圈。
    """
    if settings.teacher_marks <= 0 or not line_baselines:
        return

    mark_count = min(
        len(line_baselines),
        {
            1: rng.randint(1, 2),
            2: rng.randint(2, 4),
            3: rng.randint(4, 7),
        }.get(settings.teacher_marks, 2),
    )

    selected_baselines = rng.sample(
        line_baselines,
        k=mark_count,
    )

    draw = ImageDraw.Draw(page)
    red = (185, 46, 52, 210)

    for baseline in selected_baselines:
        style = rng.choice(
            ["check", "underline", "circle"]
        )

        if style == "check":
            x = PAGE_WIDTH - RIGHT_MARGIN + rng.randint(-14, 9)
            y = baseline - rng.randint(13, 22)

            draw.line(
                (
                    x - 8,
                    y + 8,
                    x - 2,
                    y + 14,
                    x + 12,
                    y - 4,
                ),
                fill=red,
                width=2,
                joint="curve",
            )

        elif style == "underline":
            start_x = rng.randint(
                text_start_x + 40,
                max(
                    text_start_x + 45,
                    PAGE_WIDTH - RIGHT_MARGIN - 240,
                ),
            )

            length = rng.randint(75, 180)
            y = baseline + rng.randint(3, 7)

            points = sample_quadratic_curve(
                (start_x, y),
                (
                    start_x + length // 2,
                    y + rng.randint(-3, 3),
                ),
                (
                    min(
                        PAGE_WIDTH - RIGHT_MARGIN,
                        start_x + length,
                    ),
                    y + rng.randint(-2, 2),
                ),
                steps=24,
            )

            draw.line(
                points,
                fill=red,
                width=2,
                joint="curve",
            )

        else:
            center_x = PAGE_WIDTH - RIGHT_MARGIN - rng.randint(35, 90)
            center_y = baseline - rng.randint(10, 20)
            radius_x = rng.randint(12, 20)
            radius_y = rng.randint(8, 15)

            draw.ellipse(
                (
                    center_x - radius_x,
                    center_y - radius_y,
                    center_x + radius_x,
                    center_y + radius_y,
                ),
                outline=red,
                width=2,
            )


def render_single_page(
    page_lines: list[str],
    settings: RenderSettings,
    page_number: int,
) -> Image.Image:
    page_seed = (
        settings.seed
        + page_number * 100_003
    )

    rng = random.Random(page_seed)

    image = create_paper_background(
        settings=settings,
        rng=rng,
    )

    text_start_x = get_text_start_x(
        settings
    )

    draw_notebook_header(
        page=image,
        settings=settings,
        rng=rng,
    )

    previous_was_blank = True
    written_baselines: list[int] = []

    for row_index, line in enumerate(
        page_lines
    ):
        rule_y = (
            TOP_MARGIN
            + row_index
            * settings.rule_spacing
        )

        if not line.strip():
            previous_was_blank = True
            continue

        paragraph_indent = (
            rng.randint(8, 20)
            if previous_was_blank
            else 0
        )

        x_start = (
            text_start_x
            + paragraph_indent
            + rng.randint(
                -max(1, settings.randomness),
                max(2, settings.randomness + 1),
            )
        )

        baseline_y = (
            rule_y
            - 2
            + rng.randint(
                -max(1, settings.randomness // 2),
                max(1, settings.randomness // 2),
            )
        )

        draw_handwritten_line(
            page=image,
            text=line,
            baseline_y=baseline_y,
            x_start=x_start,
            settings=settings,
            rng=rng,
        )

        written_baselines.append(
            baseline_y
        )

        previous_was_blank = False

    draw_teacher_feedback(
        page=image,
        line_baselines=written_baselines,
        text_start_x=text_start_x,
        settings=settings,
        rng=rng,
    )

    draw_page_number(
        page=image,
        page_number=page_number + 1,
        settings=settings,
        rng=rng,
    )

    return apply_page_finish(
        image,
        settings,
        rng,
    )


def render_document(
    text: str,
    settings: RenderSettings,
) -> list[Image.Image]:
    normal_font = load_font(settings.font_path, settings.font_size)

    text_start_x = get_text_start_x(
        settings
    )

    max_text_width = (
        PAGE_WIDTH
        - RIGHT_MARGIN
        - text_start_x
        - 28
    )

    lines = wrap_text(
        text=text,
        font=normal_font,
        max_width=max_text_width,
    )

    pages_lines = paginate_lines(lines=lines, settings=settings)

    return [
        render_single_page(
            page_lines=page_lines,
            settings=settings,
            page_number=index,
        )
        for index, page_lines in enumerate(pages_lines)
    ]


# ============================================================
# 导出
# ============================================================
def image_to_png_bytes(image: Image.Image) -> bytes:
    buffer = BytesIO()
    image.save(buffer, format="PNG", optimize=True)
    return buffer.getvalue()


def pages_to_zip_bytes(pages: list[Image.Image]) -> bytes:
    buffer = BytesIO()

    with zipfile.ZipFile(
        buffer,
        mode="w",
        compression=zipfile.ZIP_DEFLATED,
    ) as archive:
        for index, page in enumerate(pages, start=1):
            archive.writestr(
                f"page_{index:03d}.png",
                image_to_png_bytes(page),
            )

    return buffer.getvalue()


def pages_to_pdf_bytes(pages: list[Image.Image]) -> bytes:
    buffer = BytesIO()
    rgb_pages = [page.convert("RGB") for page in pages]

    rgb_pages[0].save(
        buffer,
        format="PDF",
        save_all=True,
        append_images=rgb_pages[1:],
        resolution=150.0,
    )

    return buffer.getvalue()


def save_result_to_session(pages: list[Image.Image]) -> None:
    st.session_state["generated_result"] = {
        "preview_pages": pages[:5],
        "page_count": len(pages),
        "first_png": image_to_png_bytes(pages[0]),
        "zip_bytes": pages_to_zip_bytes(pages),
        "pdf_bytes": pages_to_pdf_bytes(pages),
    }


# ============================================================
# Streamlit 页面
# ============================================================


st.set_page_config(
    page_title="手写生成器",
    page_icon="✒️",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        "About": "手写生成器：将 PDF、Word、TXT 或俄语文字转换为自然手写图片。",
    },
)


# ============================================================
# 页面样式
# ============================================================
st.markdown(
    """
    <style>
    :root {
        --primary-color: #536df5;
        --primary: #536df5;
        --primary-dark: #415ee6;
        --purple: #9255e9;
        --ink: #12275b;
        --muted: #7280a1;
        --border: #dfe5f3;
        --soft-border: rgba(88, 109, 182, 0.14);
        --surface: rgba(255, 255, 255, 0.96);
        --shadow: 0 14px 38px rgba(46, 66, 128, 0.09);
    }

    html {
        scroll-behavior: smooth;
    }

    .stApp {
        background:
            radial-gradient(circle at 12% 0%, rgba(71, 130, 255, 0.11), transparent 27%),
            radial-gradient(circle at 94% 2%, rgba(147, 77, 235, 0.10), transparent 28%),
            linear-gradient(180deg, #f8faff 0%, #f6f8fe 52%, #fbfcff 100%);
        color: var(--ink);
    }

    [data-testid="stHeader"],
    [data-testid="stToolbar"],
    [data-testid="stDecoration"],
    [data-testid="stStatusWidget"],
    footer,
    #MainMenu {
        display: none !important;
        visibility: hidden !important;
    }

    .block-container {
        max-width: 1510px;
        padding-top: 1.15rem;
        padding-bottom: 2.7rem;
    }

    section[data-testid="stSidebar"] {
        width: 334px !important;
        min-width: 334px !important;
        background: rgba(252, 253, 255, 0.98);
        border-right: 1px solid rgba(78, 102, 177, 0.12);
        box-shadow: 10px 0 34px rgba(44, 62, 119, 0.05);
    }

    section[data-testid="stSidebar"] > div {
        padding-top: 1rem;
        padding-left: 1.15rem;
        padding-right: 1.15rem;
    }

    section[data-testid="stSidebar"] label,
    section[data-testid="stSidebar"] p {
        color: #263a6c !important;
    }

    .sidebar-brand {
        display: flex;
        align-items: center;
        gap: 12px;
        padding: 7px 2px 16px;
        margin-bottom: 8px;
        border-bottom: 1px solid rgba(82, 105, 177, 0.12);
    }

    .sidebar-logo,
    .hero-logo {
        display: grid;
        place-items: center;
        color: #fff;
        background: linear-gradient(135deg, #4779ff, #8957ec);
        box-shadow: 0 10px 23px rgba(78, 84, 224, 0.25);
    }

    .sidebar-logo {
        width: 43px;
        height: 43px;
        border-radius: 13px;
        font-size: 22px;
    }

    .sidebar-title {
        color: #12275b;
        font-size: 19px;
        font-weight: 850;
        line-height: 1.15;
    }

    .sidebar-copy {
        margin-top: 3px;
        color: #7784a4;
        font-size: 12px;
    }

    .sidebar-label {
        margin: 14px 0 7px;
        color: #12275b;
        font-size: 15px;
        font-weight: 850;
    }

    /* 下拉框：去黑边、去黑底 */
    [data-baseweb="select"] {
        width: 100%;
    }

    div[data-baseweb="select"] > div {
        min-height: 45px !important;
        border: 1px solid var(--border) !important;
        border-radius: 13px !important;
        background: #ffffff !important;
        box-shadow: none !important;
        overflow: hidden !important;
    }

    div[data-baseweb="select"] > div:hover {
        border-color: #c3cef0 !important;
        background: #ffffff !important;
    }

    div[data-baseweb="select"] > div:focus-within {
        border-color: #7184ec !important;
        box-shadow: 0 0 0 3px rgba(83, 109, 245, 0.11) !important;
        background: #ffffff !important;
    }

    div[data-baseweb="select"] > div > div,
    div[data-baseweb="select"] > div > div > div,
    div[data-baseweb="select"] > div span,
    div[data-baseweb="select"] > div input {
        background: transparent !important;
        color: #263a6c !important;
    }

    div[data-baseweb="select"] > div > div:last-child {
        background: linear-gradient(180deg, #f7f9ff, #eef3ff) !important;
        border-left: 1px solid #e1e7f5 !important;
    }

    div[data-baseweb="select"] svg {
        fill: #6f7fa8 !important;
        color: #6f7fa8 !important;
    }

    div[role="listbox"] {
        border: 1px solid var(--border) !important;
        border-radius: 12px !important;
        background: #ffffff !important;
        box-shadow: 0 14px 36px rgba(39, 54, 103, 0.14) !important;
        overflow: hidden;
    }

    div[role="option"] {
        color: #263a6c !important;
        background: #ffffff !important;
    }

    div[role="option"]:hover,
    div[role="option"][aria-selected="true"] {
        background: #eef2ff !important;
    }

    [data-testid="stTextArea"] textarea,
    [data-testid="stTextInput"] input,
    [data-testid="stNumberInput"] input,
    input[type="text"],
    textarea {
        color: #263a6c !important;
        caret-color: #536df5 !important;
        background: #ffffff !important;
        border: 1px solid var(--border) !important;
        border-radius: 13px !important;
        outline: none !important;
        box-shadow: none !important;
    }

    [data-testid="stTextArea"] textarea:hover,
    [data-testid="stTextInput"] input:hover,
    [data-testid="stNumberInput"] input:hover {
        border-color: #c3cef0 !important;
    }

    [data-testid="stTextArea"] textarea:focus,
    [data-testid="stTextInput"] input:focus,
    [data-testid="stNumberInput"] input:focus {
        border-color: #7184ec !important;
        box-shadow: 0 0 0 3px rgba(83, 109, 245, 0.11) !important;
        outline: none !important;
    }

    button,
    button:focus,
    button:active {
        outline: none !important;
        box-shadow: none !important;
    }

    [data-testid="stFileUploaderDropzone"] {
        border: 1.5px dashed #bdc9ee !important;
        border-radius: 15px !important;
        background: linear-gradient(145deg, #f9fbff, #fbf9ff) !important;
    }

    [data-testid="stFileUploaderDropzone"]:hover {
        border-color: #7789ef !important;
        background: linear-gradient(145deg, #f5f8ff, #faf6ff) !important;
    }

    [data-testid="stFileUploaderDropzone"] button,
    [data-testid="stFileUploader"] button {
        color: #4762cb !important;
        background: linear-gradient(180deg, #f3f6ff, #eaf0ff) !important;
        border: 1px solid #d7dff6 !important;
        border-radius: 11px !important;
        box-shadow: none !important;
        font-weight: 760 !important;
    }

    [data-testid="stFileUploaderDropzone"] button:hover,
    [data-testid="stFileUploader"] button:hover {
        color: #3550bb !important;
        background: linear-gradient(180deg, #edf2ff, #e1eaff) !important;
        border-color: #c7d2f3 !important;
    }

    [data-testid="stSlider"] [role="slider"] {
        background: #ffffff !important;
        border: 2px solid #536df5 !important;
        box-shadow: 0 2px 8px rgba(63, 81, 217, 0.20) !important;
    }

    [data-testid="stCheckbox"] svg {
        color: #536df5 !important;
    }

    section[data-testid="stSidebar"] [data-testid="stExpander"] {
        margin-top: 9px;
        border: 1px solid var(--soft-border);
        border-radius: 14px;
        background: rgba(249, 250, 255, 0.92);
        overflow: hidden;
    }

    .hero-shell {
        position: relative;
        overflow: hidden;
        display: flex;
        align-items: center;
        justify-content: space-between;
        gap: 22px;
        padding: 20px 23px;
        margin-bottom: 14px;
        border: 1px solid var(--soft-border);
        border-radius: 21px;
        background: linear-gradient(118deg, rgba(255,255,255,0.98), rgba(247,248,255,0.94));
        box-shadow: var(--shadow);
    }

    .hero-shell::after {
        content: "";
        position: absolute;
        right: -80px;
        top: -110px;
        width: 250px;
        height: 250px;
        border-radius: 999px;
        background: radial-gradient(circle, rgba(145,83,235,0.14), transparent 69%);
        pointer-events: none;
    }

    .hero-brand {
        position: relative;
        z-index: 1;
        display: flex;
        align-items: center;
        gap: 15px;
    }

    .hero-logo {
        width: 54px;
        height: 54px;
        border-radius: 16px;
        font-size: 27px;
    }

    .hero-title {
        margin: 0;
        color: #10245c;
        font-size: clamp(26px, 2.1vw, 36px);
        line-height: 1.1;
        font-weight: 900;
    }

    .hero-subtitle {
        margin-top: 6px;
        color: #6b789b;
        font-size: 14px;
    }

    .hero-tip {
        position: relative;
        z-index: 1;
        padding: 10px 14px;
        border: 1px solid #dfe4f3;
        border-radius: 12px;
        color: #344c86;
        background: rgba(255,255,255,0.82);
        font-size: 13px;
        font-weight: 750;
        white-space: nowrap;
    }

    .steps-row {
        display: grid;
        grid-template-columns: repeat(3, minmax(0, 1fr));
        gap: 10px;
        margin: 0 0 16px;
    }

    .step-card {
        display: flex;
        align-items: center;
        gap: 10px;
        padding: 12px 14px;
        border: 1px solid var(--soft-border);
        border-radius: 14px;
        background: rgba(255,255,255,0.84);
    }

    .step-number {
        display: grid;
        place-items: center;
        width: 29px;
        height: 29px;
        flex: 0 0 auto;
        border-radius: 9px;
        color: #ffffff;
        background: linear-gradient(135deg, #5175f5, #8a5ae8);
        font-size: 13px;
        font-weight: 850;
    }

    .step-title {
        color: #243b70;
        font-size: 13px;
        font-weight: 800;
    }

    .step-copy {
        margin-top: 2px;
        color: #8490ad;
        font-size: 11px;
    }

    [data-testid="stVerticalBlockBorderWrapper"] {
        border: 1px solid var(--soft-border) !important;
        border-radius: 19px !important;
        background: rgba(255,255,255,0.91) !important;
        box-shadow: var(--shadow);
        overflow: hidden;
    }

    [data-testid="stVerticalBlockBorderWrapper"] > div {
        padding: 1.05rem 1.1rem 1.15rem;
    }

    .panel-header {
        display: flex;
        align-items: center;
        justify-content: space-between;
        gap: 12px;
        margin-bottom: 10px;
    }

    .panel-title {
        color: #12275b;
        font-size: 20px;
        font-weight: 880;
    }

    .panel-description {
        margin-top: 3px;
        color: #7b87a5;
        font-size: 12px;
    }

    .panel-chip {
        padding: 6px 10px;
        border-radius: 999px;
        color: #5b50ce;
        background: rgba(91,80,206,0.09);
        font-size: 11px;
        font-weight: 800;
        white-space: nowrap;
    }

    [data-testid="stTextArea"] textarea {
        min-height: 285px !important;
        line-height: 1.62;
        resize: vertical !important;
    }

    .stButton > button {
        min-height: 44px;
        border-radius: 12px !important;
        font-weight: 780 !important;
        color: #425a97 !important;
        background: linear-gradient(180deg, #f8faff, #eef3ff) !important;
        border: 1px solid #d9e1f5 !important;
        box-shadow: none !important;
    }

    .stButton > button:hover {
        color: #30498a !important;
        background: linear-gradient(180deg, #eef3ff, #e4ecff) !important;
        border-color: #c5d1f0 !important;
    }

    .stButton > button[kind="primary"] {
        min-height: 48px;
        color: #ffffff !important;
        background: linear-gradient(105deg, #4774f7, #984ce9) !important;
        border: 0 !important;
        box-shadow: 0 11px 23px rgba(75, 74, 220, 0.21) !important;
    }

    .stButton > button[kind="primary"]:hover {
        transform: translateY(-1px);
        box-shadow: 0 14px 27px rgba(75, 74, 220, 0.27) !important;
        color: #ffffff !important;
        background: linear-gradient(105deg, #406cf0, #9145e3) !important;
    }

    .stDownloadButton > button {
        min-height: 42px;
        color: #354d88 !important;
        background: linear-gradient(180deg, #fafcff, #eef3ff) !important;
        border: 1px solid #dce3f3 !important;
        border-radius: 11px !important;
        font-weight: 750 !important;
        box-shadow: none !important;
    }

    .stDownloadButton > button:hover {
        background: linear-gradient(180deg, #edf2ff, #e2eaff) !important;
        border-color: #c7d1ef !important;
    }

    [data-testid="stImage"] img {
        border-radius: 14px;
        box-shadow: 0 12px 31px rgba(41, 55, 104, 0.12);
    }

    .preview-placeholder {
        position: relative;
        display: grid;
        place-items: center;
        min-height: 560px;
        padding: 28px;
        overflow: hidden;
        border: 1px solid #e2e7f1;
        border-radius: 15px;
        background:
            linear-gradient(90deg, transparent 0 10.8%, rgba(222,107,119,0.30) 10.8% 11%, transparent 11%),
            repeating-linear-gradient(
                180deg,
                #faf8ef 0,
                #faf8ef 33px,
                rgba(108,144,211,0.26) 34px,
                #faf8ef 35px
            );
    }

    .preview-empty {
        max-width: 325px;
        padding: 20px 22px;
        border: 1px solid rgba(88,108,180,0.12);
        border-radius: 15px;
        background: rgba(255,255,255,0.88);
        box-shadow: 0 13px 34px rgba(49,65,116,0.09);
        text-align: center;
    }

    .preview-empty-icon {
        margin-bottom: 8px;
        font-size: 35px;
    }

    .preview-empty-title {
        color: #20386e;
        font-size: 16px;
        font-weight: 850;
    }

    .preview-empty-copy {
        margin-top: 6px;
        color: #7f8aa7;
        font-size: 12px;
        line-height: 1.55;
    }

    .result-bar {
        display: flex;
        align-items: center;
        justify-content: space-between;
        gap: 12px;
        margin-bottom: 11px;
        padding: 9px 12px;
        border-radius: 11px;
        color: #354e88;
        background: linear-gradient(90deg, rgba(76,113,244,0.08), rgba(148,79,230,0.07));
        font-size: 12px;
        font-weight: 800;
    }

    .footer-note {
        margin-top: 20px;
        padding: 13px 18px;
        border-radius: 15px;
        color: #ffffff;
        background: linear-gradient(105deg, #4774f7, #984ce9);
        box-shadow: 0 12px 28px rgba(74,72,211,0.18);
        text-align: center;
        font-size: 13px;
        font-weight: 800;
    }

    @media (max-width: 1050px) {
        .steps-row {
            grid-template-columns: 1fr;
        }

        .hero-tip {
            display: none;
        }

        .preview-placeholder {
            min-height: 420px;
        }
    }

    @media (max-width: 700px) {
        .block-container {
            padding-left: 0.75rem;
            padding-right: 0.75rem;
        }

        .hero-shell {
            padding: 17px;
        }

        .hero-logo {
            width: 46px;
            height: 46px;
        }

        .hero-title {
            font-size: 26px;
        }
    }
    </style>
    """,
    unsafe_allow_html=True,
)



font_files = get_font_files()

if not font_files:
    st.error("没有找到字体。请上传支持俄语的 .ttf 或 .otf 手写字体。")
    st.code(
        "app.py\n"
        "fonts/\n"
        "    BadScript-Regular.ttf"
    )
    st.stop()


# ============================================================
# 页面头部
# ============================================================
st.markdown(
    """
    <div class="hero-shell">
        <div class="hero-brand">
            <div class="hero-logo">✒</div>
            <div>
                <h1 class="hero-title">手写生成器</h1>
                <div class="hero-subtitle">
                    将 PDF、Word、TXT 或俄语文字转换为自然手写图片
                </div>
            </div>
        </div>
        <div class="hero-tip">支持多种笔迹、涂改与真实纸张效果</div>
    </div>

    <div class="steps-row">
        <div class="step-card">
            <div class="step-number">1</div>
            <div>
                <div class="step-title">输入内容</div>
                <div class="step-copy">粘贴文字或上传文档</div>
            </div>
        </div>
        <div class="step-card">
            <div class="step-number">2</div>
            <div>
                <div class="step-title">选择风格</div>
                <div class="step-copy">在左侧调整字体与纸张</div>
            </div>
        </div>
        <div class="step-card">
            <div class="step-number">3</div>
            <div>
                <div class="step-title">生成并下载</div>
                <div class="step-copy">预览后导出 PNG 或 PDF</div>
            </div>
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)


# ============================================================
# 侧边栏设置
# ============================================================
with st.sidebar:
    st.markdown(
        """
        <div class="sidebar-brand">
            <div class="sidebar-logo">✒</div>
            <div>
                <div class="sidebar-title">手写设置</div>
                <div class="sidebar-copy">常用设置在上方，高级设置可展开</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(
        '<div class="sidebar-label">常用设置</div>',
        unsafe_allow_html=True,
    )

    preset_name = st.selectbox(
        "快速预设",
        options=list(PRESETS.keys()),
        help="先选择一个接近目标的预设，再微调。",
    )

    preset = PRESETS[preset_name]
    preset_key = preset_name.replace(" ", "_")

    selected_font_name = st.selectbox(
        "手写字体",
        options=list(font_files.keys()),
        help="俄语手写推荐 Marck Script 或 Bad Script。",
        key=f"font_{preset_key}",
    )

    font_size = st.slider(
        "字体大小",
        min_value=36,
        max_value=72,
        value=int(preset["font_size"]),
        key=f"font_size_{preset_key}",
    )

    font_weight = st.slider(
        "字体粗细",
        min_value=0.75,
        max_value=1.60,
        value=1.00,
        step=0.05,
        help="1.00 为默认值。",
        key=f"font_weight_{preset_key}",
    )

    pen_options = ["中性笔", "钢笔", "圆珠笔", "铅笔"]
    pen_style = st.selectbox(
        "书写工具",
        pen_options,
        index=option_index(
            pen_options,
            preset["pen_style"],
        ),
        key=f"pen_{preset_key}",
    )

    ink_options = ["黑色", "蓝色"]
    ink_name = st.selectbox(
        "墨水颜色",
        ink_options,
        index=option_index(
            ink_options,
            preset["ink_name"],
        ),
        key=f"ink_{preset_key}",
    )

    with st.expander("笔误与涂改", expanded=False):
        correction_level = st.select_slider(
            "涂改数量",
            options=[0, 1, 2, 3],
            value=int(preset["correction_level"]),
            help="0=关闭，1=少量，2=自然，3=较多。",
            key=f"correction_{preset_key}",
        )

        correction_style_options = [
            "混合",
            "局部改错型",
            "规整型",
            "乱涂型",
        ]

        correction_style = st.selectbox(
            "涂改风格",
            correction_style_options,
            index=option_index(
                correction_style_options,
                preset.get(
                    "correction_style",
                    "混合",
                ),
            ),
            key=f"correction_style_{preset_key}",
        )

        teacher_marks = st.select_slider(
            "老师红笔批改",
            options=[0, 1, 2, 3],
            value=int(preset["teacher_marks"]),
            key=f"teacher_{preset_key}",
        )

    with st.expander("纸张效果", expanded=False):
        paper_options = [
            "俄语作业本",
            "普通横线本",
            "方格本",
            "白纸",
        ]

        paper_type = st.selectbox(
            "本子类型",
            paper_options,
            index=option_index(
                paper_options,
                preset["paper_type"],
            ),
            key=f"paper_{preset_key}",
        )

        paper_age = st.slider(
            "纸张旧化",
            min_value=0.0,
            max_value=1.0,
            value=float(preset["paper_age"]),
            step=0.05,
            key=f"paper_age_{preset_key}",
        )

        wrinkle_strength = st.slider(
            "纸张褶皱",
            min_value=0.0,
            max_value=1.0,
            value=float(
                preset.get(
                    "wrinkle_strength",
                    0.0,
                )
            ),
            step=0.05,
            key=f"wrinkle_{preset_key}",
        )

        photo_options = ["无", "扫描件", "手机拍照"]
        photo_effect = st.selectbox(
            "导出外观",
            photo_options,
            index=option_index(
                photo_options,
                preset["photo_effect"],
            ),
            key=f"photo_{preset_key}",
        )

    with st.expander("排版与页眉", expanded=False):
        line_spacing = st.slider(
            "横线间距",
            min_value=4,
            max_value=24,
            value=int(preset["line_spacing"]),
            key=f"line_spacing_{preset_key}",
        )

        word_spacing = st.slider(
            "单词间距",
            min_value=-8,
            max_value=12,
            value=int(preset["word_spacing"]),
            key=f"word_spacing_{preset_key}",
        )

        texture_strength = st.slider(
            "笔迹纹理",
            min_value=0.0,
            max_value=1.0,
            value=float(preset["texture_strength"]),
            step=0.05,
            key=f"texture_{preset_key}",
        )

        randomness = st.slider(
            "自然随机程度",
            min_value=0,
            max_value=4,
            value=int(preset["randomness"]),
            key=f"random_{preset_key}",
        )

        slant = st.slider(
            "右倾程度",
            min_value=0.00,
            max_value=0.20,
            value=float(preset["slant"]),
            step=0.01,
            key=f"slant_{preset_key}",
        )

        baseline_wave = st.slider(
            "基线起伏",
            min_value=0.0,
            max_value=5.0,
            value=float(preset["baseline_wave"]),
            step=0.2,
            key=f"baseline_{preset_key}",
        )

        header_enabled = st.checkbox(
            "显示手写页眉",
            value=True,
            key=f"header_enabled_{preset_key}",
        )

        header_date = st.text_input(
            "日期",
            value="15.01",
            key=f"header_date_{preset_key}",
        )

        header_lesson = st.text_input(
            "课题",
            value="Урок русского языка",
            key=f"header_lesson_{preset_key}",
        )

        show_page_number = st.checkbox(
            "显示页码",
            value=False,
            key=f"page_number_{preset_key}",
        )

        seed = st.number_input(
            "随机种子",
            min_value=0,
            max_value=999_999,
            value=12345,
            step=1,
            key=f"seed_{preset_key}",
        )

    # 已移除用户不需要的连笔滑块，内部沿用预设值
    connection_strength = float(
        preset["connection_strength"]
    )
    flourish_level = 0

    # 折叠区域仍会执行，但以下默认值可防止不同 Streamlit 版本产生未定义变量
    if "correction_level" not in locals():
        correction_level = int(preset["correction_level"])
    if "correction_style" not in locals():
        correction_style = str(
            preset.get("correction_style", "混合")
        )
    if "teacher_marks" not in locals():
        teacher_marks = int(preset["teacher_marks"])
    if "paper_type" not in locals():
        paper_type = str(preset["paper_type"])
    if "paper_age" not in locals():
        paper_age = float(preset["paper_age"])
    if "wrinkle_strength" not in locals():
        wrinkle_strength = float(
            preset.get("wrinkle_strength", 0.0)
        )
    if "photo_effect" not in locals():
        photo_effect = str(preset["photo_effect"])
    if "line_spacing" not in locals():
        line_spacing = int(preset["line_spacing"])
    if "word_spacing" not in locals():
        word_spacing = int(preset["word_spacing"])
    if "texture_strength" not in locals():
        texture_strength = float(preset["texture_strength"])
    if "randomness" not in locals():
        randomness = int(preset["randomness"])
    if "slant" not in locals():
        slant = float(preset["slant"])
    if "baseline_wave" not in locals():
        baseline_wave = float(preset["baseline_wave"])
    if "header_enabled" not in locals():
        header_enabled = True
    if "header_date" not in locals():
        header_date = "15.01"
    if "header_lesson" not in locals():
        header_lesson = "Урок русского языка"
    if "show_page_number" not in locals():
        show_page_number = False
    if "seed" not in locals():
        seed = 12345


# ============================================================
# 输入和预览
# ============================================================
if "editor_text" not in st.session_state:
    st.session_state["editor_text"] = (
        "Это моя семья.\n"
        "Это моя семья.\n"
        "Зовут меня Мин. Я ученик сельской школы.\n"
        "Вот мои родители. Это мой папа и моя мама.\n"
        "Мы любим нашу маленькую семью."
    )


left_column, right_column = st.columns(
    [0.40, 0.60],
    gap="large",
)


with left_column:
    with st.container(border=True):
        st.markdown(
            """
            <div class="panel-header">
                <div>
                    <div class="panel-title">1. 输入文档内容</div>
                    <div class="panel-description">
                        上传文件后仍可在文本框中检查和修改
                    </div>
                </div>
                <div class="panel-chip">PDF · DOCX · TXT</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        uploaded_file = st.file_uploader(
            "上传文档",
            type=SUPPORTED_FILE_TYPES,
            label_visibility="collapsed",
            help="支持 TXT、DOCX 和文字型 PDF。",
        )

        if uploaded_file is not None:
            file_identity = (
                uploaded_file.name,
                uploaded_file.size,
            )

            if (
                st.session_state.get(
                    "last_uploaded_file"
                )
                != file_identity
            ):
                try:
                    extracted_text = extract_uploaded_file(
                        uploaded_file
                    )
                    extracted_text = normalize_text(
                        extracted_text
                    )

                    if extracted_text:
                        st.session_state["editor_text"] = extracted_text
                        st.session_state["last_uploaded_file"] = file_identity
                        st.success("文档内容已提取，可以继续编辑。")
                    else:
                        st.warning(
                            "没有提取到文字。扫描版 PDF 暂时需要先转成文字型 PDF。"
                        )

                except Exception as error:
                    st.error(f"读取文件失败：{error}")

        text = st.text_area(
            "文档内容",
            key="editor_text",
            height=300,
            placeholder="在这里输入或粘贴俄语文字……",
        )

        clear_column, generate_column = st.columns(
            [0.32, 0.68],
            gap="small",
        )

        with clear_column:
            clear_clicked = st.button(
                "清空内容",
                use_container_width=True,
            )

        with generate_column:
            generate_clicked = st.button(
                "✦ 生成手写图片",
                type="primary",
                use_container_width=True,
            )

        if clear_clicked:
            st.session_state["editor_text"] = ""
            st.session_state.pop("last_uploaded_file", None)
            st.rerun()


if generate_clicked:
    if not text.strip():
        st.warning("请先输入或上传文字。")
    else:
        settings = RenderSettings(
            font_path=str(
                font_files[selected_font_name]
            ),
            font_size=font_size,
            line_spacing=int(line_spacing),
            paper_type=str(paper_type),
            paper_age=float(paper_age),
            wrinkle_strength=float(wrinkle_strength),
            photo_effect=str(photo_effect),
            ink_name=str(ink_name),
            pen_style=str(pen_style),
            texture_strength=float(texture_strength),
            font_weight=float(font_weight),
            randomness=int(randomness),
            seed=int(seed),
            slant=float(slant),
            word_spacing=int(word_spacing),
            connection_strength=float(connection_strength),
            baseline_wave=float(baseline_wave),
            flourish_level=0,
            correction_level=int(correction_level),
            correction_style=str(correction_style),
            teacher_marks=int(teacher_marks),
            header_enabled=bool(header_enabled),
            header_date=str(header_date),
            header_lesson=str(header_lesson),
            show_page_number=bool(show_page_number),
        )

        try:
            with st.spinner("正在生成手写页面……"):
                pages = render_document(
                    text=text,
                    settings=settings,
                )
                save_result_to_session(pages)

            st.toast(
                f"生成成功，共 {len(pages)} 页。",
                icon="✅",
            )

        except OSError:
            st.error(
                "字体无法打开。请确认字体有效并支持西里尔字母。"
            )
        except MemoryError:
            st.error(
                "文字太多，内存不足。请减少文字后分批生成。"
            )
        except Exception as error:
            st.exception(error)


result = st.session_state.get("generated_result")


with right_column:
    with st.container(border=True):
        st.markdown(
            """
            <div class="panel-header">
                <div>
                    <div class="panel-title">2. 预览与下载</div>
                    <div class="panel-description">
                        生成后可下载单页图片、全部图片或 PDF
                    </div>
                </div>
                <div class="panel-chip">实时预览</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        if result:
            preview_pages = result["preview_pages"]

            st.markdown(
                f"""
                <div class="result-bar">
                    <span>✓ 手写文档已生成</span>
                    <span>{result["page_count"]} 页</span>
                </div>
                """,
                unsafe_allow_html=True,
            )

            st.image(
                preview_pages[0],
                use_container_width=True,
            )

            download_1, download_2, download_3 = st.columns(
                3,
                gap="small",
            )

            with download_1:
                st.download_button(
                    "第一页 PNG",
                    data=result["first_png"],
                    file_name="handwriting_page_001.png",
                    mime="image/png",
                    use_container_width=True,
                )

            with download_2:
                st.download_button(
                    "全部图片 ZIP",
                    data=result["zip_bytes"],
                    file_name="handwriting_pages.zip",
                    mime="application/zip",
                    use_container_width=True,
                )

            with download_3:
                st.download_button(
                    "多页 PDF",
                    data=result["pdf_bytes"],
                    file_name="handwriting_document.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                )

        else:
            st.markdown(
                """
                <div class="preview-placeholder">
                    <div class="preview-empty">
                        <div class="preview-empty-icon">✍️</div>
                        <div class="preview-empty-title">
                            生成结果将在这里显示
                        </div>
                        <div class="preview-empty-copy">
                            输入文字并选择左侧的手写风格，
                            然后点击“生成手写图片”。
                        </div>
                    </div>
                </div>
                """,
                unsafe_allow_html=True,
            )


if result and len(result["preview_pages"]) > 1:
    with st.expander("查看其余预览页面", expanded=False):
        for page_index, page in enumerate(
            result["preview_pages"][1:],
            start=2,
        ):
            st.image(
                page,
                caption=f"第 {page_index} 页",
                use_container_width=True,
            )


st.markdown(
    """
    <div class="footer-note">
        ✦ 手写生成器 · 输入内容 → 选择风格 → 生成下载
    </div>
    """,
    unsafe_allow_html=True,
)
